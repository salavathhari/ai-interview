import os
import json
from datetime import datetime, timezone
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import func as sql_func
from io import BytesIO

from app.database import get_db
from app.auth.utils import get_current_user
from app.models.user import User
from app.models.generated_report import GeneratedReport
from app.services.report_service import ReportService
from app.core.rate_limit import limiter

router = APIRouter(
    prefix="/reports",
    tags=["reports"]
)


class ReportResponse(BaseModel):
    id: int
    title: str
    report_type: str
    status: str
    file_size: Optional[int] = None
    summary: Optional[str] = None
    scores_snapshot: Optional[str] = None
    is_outdated: bool = False
    outdated_reason: Optional[str] = None
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None

    class Config:
        from_attributes = True


class ReportListResponse(BaseModel):
    reports: List[ReportResponse]
    total: int
    page: int
    per_page: int


class ReportHistoryResponse(BaseModel):
    reports: List[ReportResponse]
    total: int


VALID_REPORT_TYPES = {"portfolio", "interview", "coding", "ats", "skill-gap", "career-readiness"}


@router.get("", response_model=ReportListResponse)
def list_reports(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    report_type: Optional[str] = None,
    include_all: bool = Query(False, description="If true, returns all reports without pagination"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all generated reports for the current user. Use include_all=true for full history."""
    query = db.query(GeneratedReport).filter(GeneratedReport.user_id == current_user.id)

    if report_type:
        if report_type not in VALID_REPORT_TYPES:
            raise HTTPException(status_code=400, detail=f"Invalid report type. Valid: {', '.join(VALID_REPORT_TYPES)}")
        query = query.filter(GeneratedReport.report_type == report_type)

    total = query.count()

    if include_all:
        reports = query.order_by(GeneratedReport.created_at.desc()).all()
        return ReportListResponse(
            reports=[ReportResponse.model_validate(r) for r in reports],
            total=total,
            page=1,
            per_page=total,
        )

    reports = query.order_by(GeneratedReport.created_at.desc()).offset(
        (page - 1) * per_page
    ).limit(per_page).all()

    return ReportListResponse(
        reports=[ReportResponse.model_validate(r) for r in reports],
        total=total,
        page=page,
        per_page=per_page,
    )


@router.get("/history", response_model=ReportHistoryResponse)
def get_report_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get full report generation history for the current user (all reports, newest first)."""
    reports = db.query(GeneratedReport).filter(
        GeneratedReport.user_id == current_user.id
    ).order_by(GeneratedReport.created_at.desc()).all()

    return ReportHistoryResponse(
        reports=[ReportResponse.model_validate(r) for r in reports],
        total=len(reports),
    )


@router.get("/{report_id}", response_model=ReportResponse)
def get_report(
    report_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get a specific report by ID."""
    report = db.query(GeneratedReport).filter(
        GeneratedReport.id == report_id,
        GeneratedReport.user_id == current_user.id,
    ).first()

    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    return ReportResponse.model_validate(report)


@router.post("/generate", response_model=ReportResponse)
def generate_report(
    report_type: str = Query("portfolio", pattern="^(portfolio|interview|coding|ats|skill-gap|career-readiness)$"),
    title: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate a new comprehensive report."""
    report_title = title or f"{report_type.replace('-', ' ').title()} Report - {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')}"

    report = GeneratedReport(
        user_id=current_user.id,
        title=report_title,
        report_type=report_type,
        status="generating",
    )
    db.add(report)
    db.commit()
    db.refresh(report)

    try:
        if report_type == "portfolio":
            pdf_buffer = ReportService.generate_comprehensive_report(db, current_user)
        elif report_type == "interview":
            pdf_buffer = ReportService.generate_interview_report(db, current_user)
        elif report_type == "coding":
            pdf_buffer = ReportService.generate_coding_report(db, current_user)
        elif report_type == "ats":
            pdf_buffer = ReportService.generate_ats_report(db, current_user)
        elif report_type == "skill-gap":
            pdf_buffer = ReportService.generate_skill_gap_report(db, current_user)
        elif report_type == "career-readiness":
            pdf_buffer = ReportService.generate_career_readiness_report(db, current_user)
        else:
            pdf_buffer = ReportService.generate_comprehensive_report(db, current_user)

        report_dir = "uploads/reports"
        os.makedirs(report_dir, exist_ok=True)
        file_path = os.path.join(report_dir, f"report_{report.id}.pdf")

        with open(file_path, "wb") as f:
            f.write(pdf_buffer.getvalue())

        file_size = os.path.getsize(file_path)
        report.file_path = file_path
        report.file_size = file_size
        report.status = "ready"

        report.scores_snapshot = json.dumps(
            ReportService._capture_scores_snapshot(db, current_user)
        )

        db.commit()
        db.refresh(report)

        return ReportResponse.model_validate(report)

    except HTTPException:
        raise
    except Exception as e:
        report.status = "failed"
        db.commit()
        raise HTTPException(status_code=500, detail="Report generation failed. Please try again later.")


@router.get("/{report_id}/download")
def download_report(
    report_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Download a generated report PDF."""
    report = db.query(GeneratedReport).filter(
        GeneratedReport.id == report_id,
        GeneratedReport.user_id == current_user.id,
    ).first()

    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    if report.status != "ready" or not report.file_path:
        raise HTTPException(status_code=400, detail="Report is not ready for download")

    if not os.path.exists(report.file_path):
        raise HTTPException(status_code=404, detail="Report file not found on disk")

    with open(report.file_path, "rb") as f:
        pdf_content = f.read()

    safe_title = report.title.replace(" ", "_").lower()
    return StreamingResponse(
        BytesIO(pdf_content),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
    )


@router.post("/{report_id}/download-docx")
def download_report_docx(
    report_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Download a generated report as DOCX."""
    report = db.query(GeneratedReport).filter(
        GeneratedReport.id == report_id,
        GeneratedReport.user_id == current_user.id,
    ).first()

    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    if report.status != "ready":
        raise HTTPException(status_code=400, detail="Report is not ready for download")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(11)
        style.font.name = 'Calibri'

        title_para = doc.add_heading(report.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {report.created_at.strftime('%Y-%m-%d %H:%M') if report.created_at else 'N/A'}")
        doc.add_paragraph("Platform: AI Career Preparation Platform")
        doc.add_paragraph("")

        scores = {}
        if report.scores_snapshot:
            try:
                scores = json.loads(report.scores_snapshot)
            except (json.JSONDecodeError, TypeError):
                pass

        if scores:
            doc.add_heading("Scores Overview", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Score"
            for key, label in [("resume_match", "Resume Match"), ("ats_score", "ATS Score"),
                               ("skill_gap_match", "Skill Gap Match"), ("overall_readiness", "Career Readiness"),
                               ("interview_score", "Interview Score"), ("coding_score", "Coding Score")]:
                if key in scores:
                    row = table.add_row().cells
                    row[0].text = label
                    row[1].text = f"{scores[key]:.1f}%"
            doc.add_paragraph("")

        if report.summary:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(report.summary)

        if report.outdated_reason:
            doc.add_paragraph("")
            note = doc.add_paragraph()
            run = note.add_run(f"Note: This report may be outdated - {report.outdated_reason}")
            run.font.color.rgb = RGBColor(217, 119, 6)

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        safe_title = report.title.replace(" ", "_").lower()
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="DOCX export not available (python-docx not installed)")
    except Exception as e:
        raise HTTPException(status_code=500, detail="DOCX export failed. Please try again later.")


@router.delete("/{report_id}")
def delete_report(
    report_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete a generated report."""
    report = db.query(GeneratedReport).filter(
        GeneratedReport.id == report_id,
        GeneratedReport.user_id == current_user.id,
    ).first()

    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    if report.file_path and os.path.exists(report.file_path):
        try:
            os.remove(report.file_path)
        except OSError:
            pass

    db.delete(report)
    db.commit()

    return {"message": "Report deleted successfully"}


@router.post("/{report_id}/regenerate", response_model=ReportResponse)
def regenerate_report(
    report_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Regenerate an existing report (creates a new version)."""
    existing_report = db.query(GeneratedReport).filter(
        GeneratedReport.id == report_id,
        GeneratedReport.user_id == current_user.id,
    ).first()

    if not existing_report:
        raise HTTPException(status_code=404, detail="Report not found")

    return generate_report(
        report_type=existing_report.report_type,
        title=f"{existing_report.title} (Updated)",
        current_user=current_user,
        db=db,
    )


@router.post("/mark-outdated")
def mark_reports_outdated(
    reason: str = Query(..., description="Reason for marking reports as outdated"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Mark all reports as outdated (called when underlying data changes)."""
    updated = db.query(GeneratedReport).filter(
        GeneratedReport.user_id == current_user.id,
        GeneratedReport.is_outdated == False,
    ).update({
        "is_outdated": True,
        "outdated_reason": reason,
    })
    db.commit()

    return {"message": f"Marked {updated} reports as outdated", "count": updated}
