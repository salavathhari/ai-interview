import os
import json
import hashlib
import shutil
import io
import uuid
import pdfplumber
from pathlib import Path
from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, status, Query, Request, Body
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func as sa_func
from typing import List, Optional
from pydantic import BaseModel
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from app.database import get_db
from app.auth.utils import get_current_user
from app.core.rate_limit import limiter
from app.models.user import User
from app.models.resume import Resume
from app.models.interview_session import InterviewSession
from app.models.coding_challenge import CodingSession
from app.models.career import (
    JobDescription,
    ResumeAnalysis,
    SkillGapAnalysis,
    LearningRoadmap,
    OptimizedResume,
    CareerReadiness,
    CareerReadinessHistory,
)
from app.schemas.career import (
    JobDescriptionResponse,
    JobDescriptionAnalysisResponse,
    ResumeAnalysisResponse,
    SkillGapResponse,
    LearningRoadmapResponse,
    OptimizedResumeResponse,
    CareerReadinessResponse,
    CareerReadinessHistoryResponse,
    CareerDashboardResponse,
)
from app.services.career_service import CareerService

router = APIRouter(prefix="/career", tags=["Career Readiness"])

UPLOAD_DIR = Path("uploads/job_descriptions")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

SKILL_NORMALIZATION = {
    "node": "node.js", "nodejs": "node.js", "node js": "node.js",
    "react.js": "react", "reactjs": "react", "react js": "react",
    "vue.js": "vue", "vuejs": "vue", "vue js": "vue",
    "angular.js": "angular", "angularjs": "angular",
    "next.js": "nextjs", "nextjs": "nextjs", "next js": "nextjs",
    "express.js": "express", "expressjs": "express",
    "vuejs": "vue", "typescript": "typescript", "ts": "typescript",
    "postgres": "postgresql", "postgres": "postgresql",
    "k8s": "kubernetes",
    "tf": "terraform",
    "js": "javascript",
    "py": "python",
    "ci/cd": "ci_cd",
    "rest api": "rest_api", "rest": "rest_api",
    "object oriented programming": "oop",
    "object-oriented programming": "oop",
    "ml": "machine_learning",
    "deep learning": "deep_learning",
    "data structures and algorithms": "dsa",
    "data structures & algorithms": "dsa",
    "ds&a": "dsa",
    "git hub": "github",
    "github actions": "github_actions",
    "amazon web services": "aws",
    "google cloud platform": "gcp",
    "google cloud": "gcp",
    "microsoft azure": "azure",
}


def _normalize_skill(skill: str) -> str:
    s = skill.lower().strip()
    return SKILL_NORMALIZATION.get(s, s)


def _compute_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _is_encrypted_pdf(file_path: str) -> bool:
    try:
        with open(file_path, "rb") as f:
            header = f.read(1024)
            return b"/Encrypt" in header and b"/Standard" in header
    except Exception:
        return False


# ──────────────────────────────────────────────
#  JOB DESCRIPTIONS
# ──────────────────────────────────────────────

@router.post("/job-description/upload")
async def upload_job_description(
    title: str = Form(...),
    company: str = Form(None),
    file: UploadFile = File(None),
    raw_text: str = Form(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not file and not raw_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Either a file or raw_text must be provided.",
        )

    if title and len(title.strip()) < 1:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Title cannot be empty.",
        )

    extracted_text = ""
    file_path = None
    source = "paste"

    if file:
        allowed_exts = (".pdf", ".docx", ".doc", ".txt")
        filename_lower = file.filename.lower()
        if not filename_lower.endswith(allowed_exts):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_exts)}",
            )

        if filename_lower.endswith(".pdf"):
            source = "pdf"
        elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
            source = "docx"
        else:
            source = "txt"

        safe_name = os.path.basename(file.filename)
        uid = uuid.uuid4().hex[:12]
        stored_name = f"{current_user.id}_{uid}_{safe_name}"
        file_path = str(UPLOAD_DIR / stored_name)

        contents = await file.read()
        if len(contents) == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Uploaded file is empty.",
            )
        if len(contents) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB.",
            )

        try:
            with open(file_path, "wb") as buffer:
                buffer.write(contents)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Could not save file: {str(e)}",
            )

        if source == "pdf":
            if _is_encrypted_pdf(file_path):
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Encrypted or password-protected PDFs are not supported.",
                )
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        extracted_text += page.extract_text() or ""
                if not extracted_text.strip():
                    raise ValueError("No text could be extracted from the PDF.")
            except HTTPException:
                raise
            except Exception as e:
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Could not extract text from PDF: {str(e)}",
                )
        elif source == "docx":
            try:
                from docx import Document
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    extracted_text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        extracted_text += row_text + "\n"
                if not extracted_text.strip():
                    raise ValueError("No text could be extracted from the DOCX file.")
            except HTTPException:
                raise
            except Exception as e:
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Could not extract text from DOCX: {str(e)}",
                )
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted_text = f.read()
                if not extracted_text.strip():
                    raise ValueError("No text could be extracted from the file.")
            except HTTPException:
                raise
            except Exception as e:
                if os.path.exists(file_path):
                    os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Could not extract text from file: {str(e)}",
                )
    else:
        extracted_text = raw_text
        source = "paste"

    content_hash = _compute_hash(extracted_text)

    existing = db.query(JobDescription).filter(
        JobDescription.user_id == current_user.id,
        JobDescription.content_hash == content_hash,
    ).first()
    if existing:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"A job description with identical content already exists (ID: {existing.id}).",
        )

    jd = JobDescription(
        user_id=current_user.id,
        title=title.strip(),
        company=company.strip() if company else None,
        raw_text=extracted_text,
        source=source,
        file_path=file_path,
        content_hash=content_hash,
    )
    db.add(jd)
    db.commit()
    db.refresh(jd)

    analysis = None
    try:
        analysis = CareerService.analyze_job_description(extracted_text)
        normalized_required = [_normalize_skill(s) for s in analysis.get("required_skills", [])]
        normalized_preferred = [_normalize_skill(s) for s in analysis.get("preferred_skills", [])]
        normalized_tech = [_normalize_skill(s) for s in analysis.get("technologies", [])]
        normalized_soft = [_normalize_skill(s) for s in analysis.get("soft_skills", [])]
        normalized_kw = [_normalize_skill(s) for s in analysis.get("keywords", [])]

        seen = set()
        dedup = lambda lst: [x for x in lst if not (x in seen or seen.add(x))]

        jd.required_skills = json.dumps(dedup(normalized_required))
        jd.preferred_skills = json.dumps(dedup(normalized_preferred))
        jd.technologies = json.dumps(dedup(normalized_tech))
        jd.responsibilities = json.dumps(analysis.get("responsibilities", []))
        jd.experience_years = analysis.get("experience_years")
        jd.education_requirements = analysis.get("education_requirements")
        jd.soft_skills = json.dumps(dedup(normalized_soft))
        jd.keywords = json.dumps(dedup(normalized_kw))
        jd.is_analyzed = True
        db.commit()
        db.refresh(jd)

        try:
            from app.services.automation_service import AutomationService
            automation = AutomationService(db)
            automation.on_jd_change(current_user.id, jd.id)
        except Exception as auto_err:
            print(f"Automation trigger (JD change) failed: {auto_err}")
    except Exception as ai_err:
        print(f"JD auto-analysis failed: {ai_err}")

    return {
        "id": jd.id,
        "title": jd.title,
        "company": jd.company,
        "source": jd.source,
        "is_analyzed": jd.is_analyzed,
        "required_skills": json.loads(jd.required_skills) if jd.required_skills else [],
        "preferred_skills": json.loads(jd.preferred_skills) if jd.preferred_skills else [],
        "technologies": json.loads(jd.technologies) if jd.technologies else [],
        "soft_skills": json.loads(jd.soft_skills) if jd.soft_skills else [],
        "keywords": json.loads(jd.keywords) if jd.keywords else [],
        "experience_years": jd.experience_years,
        "education_requirements": jd.education_requirements,
        "responsibilities": json.loads(jd.responsibilities) if jd.responsibilities else [],
        "created_at": jd.created_at,
        "message": "Job description uploaded and analyzed successfully.",
    }


@router.post("/job-description/analyze/{jd_id}")
@limiter.limit("5/minute")
def analyze_job_description(
    request: Request,
    jd_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    jd = db.query(JobDescription).filter(
        JobDescription.id == jd_id,
        JobDescription.user_id == current_user.id,
    ).first()
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")

    analysis = CareerService.analyze_job_description(jd.raw_text)

    normalized_required = [_normalize_skill(s) for s in analysis.get("required_skills", [])]
    normalized_preferred = [_normalize_skill(s) for s in analysis.get("preferred_skills", [])]
    normalized_tech = [_normalize_skill(s) for s in analysis.get("technologies", [])]
    normalized_soft = [_normalize_skill(s) for s in analysis.get("soft_skills", [])]
    normalized_kw = [_normalize_skill(s) for s in analysis.get("keywords", [])]

    seen = set()
    dedup = lambda lst: [x for x in lst if not (x in seen or seen.add(x))]

    jd.required_skills = json.dumps(dedup(normalized_required))
    jd.preferred_skills = json.dumps(dedup(normalized_preferred))
    jd.technologies = json.dumps(dedup(normalized_tech))
    jd.responsibilities = json.dumps(analysis.get("responsibilities", []))
    jd.experience_years = analysis.get("experience_years")
    jd.education_requirements = analysis.get("education_requirements")
    jd.soft_skills = json.dumps(dedup(normalized_soft))
    jd.keywords = json.dumps(dedup(normalized_kw))
    jd.is_analyzed = True

    db.commit()
    db.refresh(jd)

    try:
        from app.services.automation_service import AutomationService
        automation = AutomationService(db)
        automation.on_jd_change(current_user.id, jd.id)
    except Exception as auto_err:
        print(f"Automation trigger (JD change) failed: {auto_err}")

    return {
        "id": jd.id,
        "title": jd.title,
        "company": jd.company,
        "source": jd.source,
        "is_analyzed": jd.is_analyzed,
        "required_skills": json.loads(jd.required_skills) if jd.required_skills else [],
        "preferred_skills": json.loads(jd.preferred_skills) if jd.preferred_skills else [],
        "technologies": json.loads(jd.technologies) if jd.technologies else [],
        "soft_skills": json.loads(jd.soft_skills) if jd.soft_skills else [],
        "keywords": json.loads(jd.keywords) if jd.keywords else [],
        "experience_years": jd.experience_years,
        "education_requirements": jd.education_requirements,
        "responsibilities": json.loads(jd.responsibilities) if jd.responsibilities else [],
        "created_at": jd.created_at,
    }


@router.get("/job-descriptions")
def get_job_descriptions(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    search: Optional[str] = Query(None),
):
    query = db.query(JobDescription).filter(JobDescription.user_id == current_user.id)
    if search:
        pattern = f"%{search}%"
        query = query.filter(
            (JobDescription.title.ilike(pattern)) |
            (JobDescription.company.ilike(pattern)) |
            (JobDescription.raw_text.ilike(pattern))
        )
    total = query.count()
    items = query.order_by(JobDescription.created_at.desc()).offset(
        (page - 1) * per_page
    ).limit(per_page).all()
    return {
        "items": [
            {
                "id": jd.id,
                "title": jd.title,
                "company": jd.company,
                "source": jd.source,
                "required_skills": json.loads(jd.required_skills) if jd.required_skills else [],
                "preferred_skills": json.loads(jd.preferred_skills) if jd.preferred_skills else [],
                "technologies": json.loads(jd.technologies) if jd.technologies else [],
                "experience_years": jd.experience_years,
                "education_requirements": jd.education_requirements,
                "soft_skills": json.loads(jd.soft_skills) if jd.soft_skills else [],
                "keywords": json.loads(jd.keywords) if jd.keywords else [],
                "responsibilities": json.loads(jd.responsibilities) if jd.responsibilities else [],
                "is_analyzed": jd.is_analyzed,
                "created_at": jd.created_at,
            }
            for jd in items
        ],
        "total": total,
        "page": page,
        "per_page": per_page,
        "pages": (total + per_page - 1) // per_page,
    }


@router.get("/job-description/{jd_id}")
def get_job_description(
    jd_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    jd = db.query(JobDescription).filter(
        JobDescription.id == jd_id,
        JobDescription.user_id == current_user.id,
    ).first()
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")
    return {
        "id": jd.id,
        "title": jd.title,
        "company": jd.company,
        "source": jd.source,
        "raw_text": jd.raw_text,
        "required_skills": json.loads(jd.required_skills) if jd.required_skills else [],
        "preferred_skills": json.loads(jd.preferred_skills) if jd.preferred_skills else [],
        "technologies": json.loads(jd.technologies) if jd.technologies else [],
        "responsibilities": json.loads(jd.responsibilities) if jd.responsibilities else [],
        "experience_years": jd.experience_years,
        "education_requirements": jd.education_requirements,
        "soft_skills": json.loads(jd.soft_skills) if jd.soft_skills else [],
        "keywords": json.loads(jd.keywords) if jd.keywords else [],
        "is_analyzed": jd.is_analyzed,
        "created_at": jd.created_at,
    }


@router.put("/job-description/{jd_id}")
async def update_job_description(
    jd_id: int,
    title: str = Form(None),
    company: str = Form(None),
    file: UploadFile = File(None),
    raw_text: str = Form(None),
    reanalyze: bool = Form(False),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    jd = db.query(JobDescription).filter(
        JobDescription.id == jd_id,
        JobDescription.user_id == current_user.id,
    ).first()
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")

    if title is not None:
        jd.title = title.strip()
    if company is not None:
        jd.company = company.strip() if company.strip() else None

    text_changed = False
    if file:
        allowed_exts = (".pdf", ".docx", ".doc", ".txt")
        filename_lower = file.filename.lower()
        if not filename_lower.endswith(allowed_exts):
            raise HTTPException(status_code=400, detail=f"Unsupported file type.")

        safe_name = os.path.basename(file.filename)
        uid = uuid.uuid4().hex[:12]
        stored_name = f"{current_user.id}_{uid}_{safe_name}"
        new_path = str(UPLOAD_DIR / stored_name)

        contents = await file.read()
        if len(contents) > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail="File too large.")
        if len(contents) == 0:
            raise HTTPException(status_code=400, detail="File is empty.")

        with open(new_path, "wb") as f:
            f.write(contents)

        if filename_lower.endswith(".pdf") and _is_encrypted_pdf(new_path):
            os.remove(new_path)
            raise HTTPException(status_code=422, detail="Encrypted PDF not supported.")

        new_text = ""
        if filename_lower.endswith(".pdf"):
            with pdfplumber.open(new_path) as pdf:
                for page in pdf.pages:
                    new_text += page.extract_text() or ""
        elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
            from docx import Document
            doc = Document(new_path)
            for p in doc.paragraphs:
                new_text += p.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    new_text += " | ".join(c.text for c in row.cells) + "\n"
        else:
            with open(new_path, "r", encoding="utf-8", errors="ignore") as f:
                new_text = f.read()

        if not new_text.strip():
            os.remove(new_path)
            raise HTTPException(status_code=422, detail="No text extracted.")

        if jd.file_path and os.path.exists(jd.file_path):
            os.remove(jd.file_path)

        jd.file_path = new_path
        jd.raw_text = new_text
        jd.source = "pdf" if filename_lower.endswith(".pdf") else "docx" if filename_lower.endswith((".docx", ".doc")) else "txt"
        jd.content_hash = _compute_hash(new_text)
        text_changed = True
    elif raw_text is not None:
        jd.raw_text = raw_text
        jd.source = "paste"
        jd.content_hash = _compute_hash(raw_text)
        text_changed = True

    if reanalyze or text_changed:
        try:
            analysis = CareerService.analyze_job_description(jd.raw_text)
            seen = set()
            dedup = lambda lst: [x for x in lst if not (x in seen or seen.add(x))]
            jd.required_skills = json.dumps(dedup([_normalize_skill(s) for s in analysis.get("required_skills", [])]))
            jd.preferred_skills = json.dumps(dedup([_normalize_skill(s) for s in analysis.get("preferred_skills", [])]))
            jd.technologies = json.dumps(dedup([_normalize_skill(s) for s in analysis.get("technologies", [])]))
            jd.responsibilities = json.dumps(analysis.get("responsibilities", []))
            jd.experience_years = analysis.get("experience_years")
            jd.education_requirements = analysis.get("education_requirements")
            jd.soft_skills = json.dumps(dedup([_normalize_skill(s) for s in analysis.get("soft_skills", [])]))
            jd.keywords = json.dumps(dedup([_normalize_skill(s) for s in analysis.get("keywords", [])]))
            jd.is_analyzed = True
        except Exception as e:
            print(f"JD re-analysis failed: {e}")

    db.commit()
    db.refresh(jd)

    if text_changed:
        try:
            from app.services.automation_service import AutomationService
            automation = AutomationService(db)
            automation.on_jd_change(current_user.id, jd.id)
        except Exception:
            pass

    return {
        "id": jd.id,
        "title": jd.title,
        "company": jd.company,
        "source": jd.source,
        "is_analyzed": jd.is_analyzed,
        "required_skills": json.loads(jd.required_skills) if jd.required_skills else [],
        "preferred_skills": json.loads(jd.preferred_skills) if jd.preferred_skills else [],
        "technologies": json.loads(jd.technologies) if jd.technologies else [],
        "message": "Job description updated.",
    }


# ──────────────────────────────────────────────
#  CAREER ANALYSIS
# ──────────────────────────────────────────────

@router.post("/analyze", response_model=dict)
@limiter.limit("5/minute")
def analyze_career(
    request: Request,
    resume_id: int,
    job_description_id: int = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    jd = None
    if job_description_id:
        jd = db.query(JobDescription).filter(
            JobDescription.id == job_description_id,
            JobDescription.user_id == current_user.id,
        ).first()
        if not jd:
            raise HTTPException(status_code=404, detail="Job description not found")

    resume_analysis_data = CareerService.analyze_resume(
        resume.extracted_text,
        jd.raw_text if jd else None,
    )

    ats_data = CareerService.calculate_comprehensive_ats_score(
        resume.extracted_text,
        jd.raw_text if jd else None,
    )

    analysis = ResumeAnalysis(
        user_id=current_user.id,
        resume_id=resume.id,
        job_description_id=jd.id if jd else None,
        summary=resume_analysis_data.get("summary"),
        detected_skills=json.dumps(resume_analysis_data.get("detected_skills", [])),
        experience_level=resume_analysis_data.get("experience_level"),
        projects=json.dumps(resume_analysis_data.get("projects", [])),
        technologies=json.dumps(resume_analysis_data.get("technologies", [])),
        education=json.dumps(resume_analysis_data.get("education", [])),
        certifications=json.dumps(resume_analysis_data.get("certifications", [])),
        ats_score=ats_data.get("overall_score"),
        ats_breakdown=json.dumps(ats_data.get("breakdown", {})),
        ats_suggestions=json.dumps([r["message"] for r in ats_data.get("recommendations", [])]),
        is_analyzed=True,
    )

    match_score = None
    match_breakdown = {}
    skill_gap_result = None

    if jd:
        required_skills = json.loads(jd.required_skills) if jd.required_skills else []
        preferred_skills = json.loads(jd.preferred_skills) if jd.preferred_skills else []
        tech_skills = json.loads(jd.technologies) if jd.technologies else []

        if not required_skills and jd.raw_text:
            try:
                analysis_jd = CareerService.analyze_job_description(jd.raw_text)
                required_skills = analysis_jd.get("required_skills", [])
                preferred_skills = analysis_jd.get("preferred_skills", [])
                tech_skills = analysis_jd.get("technologies", [])
                from app.services.career_service import _normalize_skill
                required_skills = [_normalize_skill(s) for s in required_skills]
                preferred_skills = [_normalize_skill(s) for s in preferred_skills]
                jd.required_skills = json.dumps(required_skills)
                jd.preferred_skills = json.dumps(preferred_skills)
                jd.technologies = json.dumps(tech_skills)
                jd.is_analyzed = True
                db.commit()
            except Exception as jd_err:
                print(f"JD re-analysis failed: {jd_err}")

        all_jd_skills = list(set(required_skills + preferred_skills + tech_skills))
        if not all_jd_skills:
            all_jd_skills = required_skills

        match_data = CareerService.analyze_skill_gap(
            resume_analysis_data.get("detected_skills", []),
            all_jd_skills,
        )
        skill_gap_result = match_data
        match_score = match_data.get("match_percentage")
        match_breakdown = {
            "matched": match_data.get("matched_skills", []),
            "missing": match_data.get("missing_skills", []),
            "additional": match_data.get("additional_skills", []),
        }

        analysis.resume_match_score = match_score
        analysis.match_breakdown = json.dumps(match_breakdown)

    db.add(analysis)
    db.commit()
    db.refresh(analysis)

    result = {
        "analysis_id": analysis.id,
        "resume_analysis": ResumeAnalysisResponse.model_validate(analysis).model_dump(),
        "ats_score": ats_data.get("overall_score"),
        "ats_breakdown": ats_data.get("breakdown"),
        "ats_suggestions": [r["message"] for r in ats_data.get("recommendations", [])],
    }

    if skill_gap_result:
        gap = SkillGapAnalysis(
            user_id=current_user.id,
            resume_analysis_id=analysis.id,
            job_description_id=jd.id,
            matched_skills=json.dumps(skill_gap_result.get("matched_skills", [])),
            missing_skills=json.dumps(skill_gap_result.get("missing_skills", [])),
            additional_skills=json.dumps(skill_gap_result.get("additional_skills", [])),
            priority_skills=json.dumps(skill_gap_result.get("priority_skills", [])),
            match_percentage=skill_gap_result.get("match_percentage"),
        )
        db.add(gap)
        db.commit()
        db.refresh(gap)

        result["skill_gap"] = SkillGapResponse.model_validate(gap).model_dump()
        result["match_score"] = match_score
        result["match_breakdown"] = match_breakdown

    # Trigger automation: update analytics and recommendations
    try:
        from app.services.automation_service import AutomationService
        automation = AutomationService(db)
        automation.on_resume_change(current_user.id, analysis.id)
    except Exception as auto_err:
        print(f"Automation trigger (resume change) failed: {auto_err}")

    return result


@router.get("/resume-analysis/{analysis_id}", response_model=ResumeAnalysisResponse)
def get_resume_analysis(
    analysis_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    analysis = db.query(ResumeAnalysis).filter(
        ResumeAnalysis.id == analysis_id,
        ResumeAnalysis.user_id == current_user.id,
    ).first()
    if not analysis:
        raise HTTPException(status_code=404, detail="Resume analysis not found")
    return analysis


@router.get("/resume-analyses", response_model=List[ResumeAnalysisResponse])
def get_resume_analyses(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id)
        .order_by(ResumeAnalysis.created_at.desc())
        .all()
    )


@router.get("/skill-gap-analyses")
def get_skill_gap_analyses(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    search: str = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all skill gap analyses for the current user with pagination."""
    query = db.query(SkillGapAnalysis).filter(
        SkillGapAnalysis.user_id == current_user.id,
    )

    if search:
        query = query.filter(
            SkillGapAnalysis.matched_skills.ilike(f"%{search}%")
            | SkillGapAnalysis.missing_skills.ilike(f"%{search}%")
            | SkillGapAnalysis.priority_skills.ilike(f"%{search}%")
        )

    total = query.count()
    analyses = (
        query.order_by(SkillGapAnalysis.created_at.desc())
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )

    return {
        "items": analyses,
        "total": total,
        "page": page,
        "per_page": per_page,
        "pages": (total + per_page - 1) // per_page,
    }


@router.get("/skill-gap/{gap_id}", response_model=SkillGapResponse)
def get_skill_gap(
    gap_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    gap = db.query(SkillGapAnalysis).filter(
        SkillGapAnalysis.id == gap_id,
        SkillGapAnalysis.user_id == current_user.id,
    ).first()
    if not gap:
        raise HTTPException(status_code=404, detail="Skill gap analysis not found")
    return gap


# ──────────────────────────────────────────────
#  LEARNING ROADMAP
# ──────────────────────────────────────────────


def _calculate_learning_streak(db: Session, user_id: int) -> int:
    """Calculate consecutive-day learning streak from LearningProgress records."""
    from datetime import date, timedelta
    from app.models.intelligence import LearningProgress

    dates = set()
    for p in db.query(LearningProgress.completed_at).filter(
        LearningProgress.user_id == user_id,
        LearningProgress.completed_at.isnot(None),
    ).all():
        if p[0]:
            dates.add(p[0].date() if hasattr(p[0], 'date') else p[0])

    streak = 0
    # Start from today; if no activity today, start from yesterday
    d = date.today()
    if d not in dates:
        d -= timedelta(days=1)
    while d in dates:
        streak += 1
        d -= timedelta(days=1)
    return streak

@router.post("/roadmap/generate", response_model=LearningRoadmapResponse)
@limiter.limit("3/minute")
def generate_roadmap(
    request: Request,
    skill_gap_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    gap = db.query(SkillGapAnalysis).filter(
        SkillGapAnalysis.id == skill_gap_id,
        SkillGapAnalysis.user_id == current_user.id,
    ).first()
    if not gap:
        raise HTTPException(status_code=404, detail="Skill gap analysis not found")

    missing = json.loads(gap.missing_skills) if gap.missing_skills else []
    roadmap_data = CareerService.generate_learning_roadmap(missing)

    roadmap = LearningRoadmap(
        user_id=current_user.id,
        skill_gap_id=gap.id,
        roadmap_items=json.dumps(roadmap_data.get("items", [])),
        total_hours=roadmap_data.get("total_hours"),
        estimated_weeks=roadmap_data.get("estimated_weeks"),
        status="active",
    )
    db.add(roadmap)
    db.commit()
    db.refresh(roadmap)
    return roadmap


@router.get("/roadmap/current")
def get_current_roadmap(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get the current active roadmap for the user."""
    roadmap = db.query(LearningRoadmap).filter(
        LearningRoadmap.user_id == current_user.id,
        LearningRoadmap.status == "active",
    ).order_by(LearningRoadmap.created_at.desc()).first()

    if not roadmap:
        return {"roadmap": None, "message": "No active roadmap. Generate one to get started."}

    return {
        "roadmap": {
            "id": roadmap.id,
            "career_goal": roadmap.career_goal,
            "total_hours": roadmap.total_hours,
            "estimated_weeks": roadmap.estimated_weeks,
            "current_readiness": roadmap.current_readiness,
            "target_readiness": roadmap.target_readiness,
            "interview_readiness": roadmap.interview_readiness,
            "coding_readiness": roadmap.coding_readiness,
            "completed_topics": roadmap.completed_topics,
            "status": roadmap.status,
            "progress_percentage": roadmap.progress_percentage,
            "version": roadmap.version,
            "created_at": str(roadmap.created_at) if roadmap.created_at else None,
        },
        "phases": json.loads(roadmap.phases) if roadmap.phases else [],
        "daily_plan": {**(json.loads(roadmap.daily_plan) if roadmap.daily_plan else {}), "streak_days": _calculate_learning_streak(db, current_user.id)},
        "mentor_tips": json.loads(roadmap.mentor_tips) if roadmap.mentor_tips else [],
        "skill_gap_summary": json.loads(roadmap.skill_gap_summary) if roadmap.skill_gap_summary else {},
    }


@router.get("/roadmap/history")
def get_roadmap_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get all roadmaps for the user (history)."""
    roadmaps = db.query(LearningRoadmap).filter(
        LearningRoadmap.user_id == current_user.id,
    ).order_by(LearningRoadmap.created_at.desc()).all()

    return {
        "roadmaps": [
            {
                "id": r.id,
                "career_goal": r.career_goal,
                "status": r.status,
                "progress_percentage": r.progress_percentage,
                "total_hours": r.total_hours,
                "estimated_weeks": r.estimated_weeks,
                "version": r.version,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in roadmaps
        ],
        "total": len(roadmaps),
    }


@router.get("/roadmap/{roadmap_id}", response_model=LearningRoadmapResponse)
def get_roadmap(
    roadmap_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    roadmap = db.query(LearningRoadmap).filter(
        LearningRoadmap.id == roadmap_id,
        LearningRoadmap.user_id == current_user.id,
    ).first()
    if not roadmap:
        raise HTTPException(status_code=404, detail="Learning roadmap not found")
    return roadmap


@router.get("/roadmaps", response_model=List[LearningRoadmapResponse])
def get_roadmaps(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return (
        db.query(LearningRoadmap)
        .filter(LearningRoadmap.user_id == current_user.id)
        .order_by(LearningRoadmap.created_at.desc())
        .all()
    )


# ──────────────────────────────────────────────
#  ENHANCED AI LEARNING PLANNER
# ──────────────────────────────────────────────

@router.post("/roadmap/generate-enhanced")
@limiter.limit("3/minute")
def generate_enhanced_roadmap(
    request: Request,
    skill_gap_id: int,
    resume_analysis_id: int = None,
    job_description_id: int = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate a production-grade learning roadmap with phases, resources, and AI mentor tips."""
    from app.services.enhanced_roadmap_generator import EnhancedRoadmapGenerator

    gap = db.query(SkillGapAnalysis).filter(
        SkillGapAnalysis.id == skill_gap_id,
        SkillGapAnalysis.user_id == current_user.id,
    ).first()
    if not gap:
        raise HTTPException(status_code=404, detail="Skill gap analysis not found")

    generator = EnhancedRoadmapGenerator(db)
    roadmap_data = generator.generate(
        user_id=current_user.id,
        skill_gap_id=skill_gap_id,
        resume_analysis_id=resume_analysis_id,
        job_description_id=job_description_id,
    )

    # Mark previous active roadmaps as archived
    db.query(LearningRoadmap).filter(
        LearningRoadmap.user_id == current_user.id,
        LearningRoadmap.status == "active",
    ).update({"status": "archived"})
    db.flush()

    # Save to database with ALL metadata
    roadmap = LearningRoadmap(
        user_id=current_user.id,
        skill_gap_id=skill_gap_id,
        roadmap_items=json.dumps(roadmap_data.get("phases", [])),
        phases=json.dumps(roadmap_data.get("phases", [])),
        current_phase_index=0,
        daily_plan=json.dumps(roadmap_data.get("daily_plan", {})),
        mentor_tips=json.dumps(roadmap_data.get("mentor_tips", [])),
        skill_gap_summary=json.dumps(roadmap_data.get("skill_gap_summary", {})),
        total_hours=roadmap_data.get("total_hours"),
        estimated_weeks=roadmap_data.get("estimated_weeks"),
        status="active",
        career_goal=roadmap_data.get("career_goal", ""),
        target_role=gap.job_description.title if gap.job_description else None,
        target_company=gap.job_description.company if gap.job_description else None,
        current_readiness=roadmap_data.get("current_readiness", 0),
        target_readiness=roadmap_data.get("target_readiness", 85),
        interview_readiness=roadmap_data.get("interview_readiness", 0),
        coding_readiness=roadmap_data.get("coding_readiness", 0),
        version=1,
    )
    db.add(roadmap)
    db.commit()
    db.refresh(roadmap)

    roadmap_data["roadmap_id"] = roadmap.id
    roadmap_data.setdefault("daily_plan", {})["streak_days"] = _calculate_learning_streak(db, current_user.id)
    return roadmap_data


@router.get("/roadmap/{roadmap_id}/analytics")
def get_roadmap_analytics(
    roadmap_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get detailed analytics for a roadmap."""
    roadmap = db.query(LearningRoadmap).filter(
        LearningRoadmap.id == roadmap_id,
        LearningRoadmap.user_id == current_user.id,
    ).first()
    if not roadmap:
        raise HTTPException(status_code=404, detail="Roadmap not found")

    phases = json.loads(roadmap.roadmap_items) if roadmap.roadmap_items else []
    completed = json.loads(roadmap.completed_topics) if roadmap.completed_topics else []

    total_topics = sum(len(p.get("topics", [])) for p in phases)
    completed_count = len(completed)

    return {
        "roadmap_id": roadmap.id,
        "total_phases": len(phases),
        "total_topics": total_topics,
        "completed_topics": completed_count,
        "progress_percentage": roadmap.progress_percentage,
        "total_hours": roadmap.total_hours,
        "estimated_weeks": roadmap.estimated_weeks,
        "status": roadmap.status,
        "phases": [{
            "phase_number": p.get("phase_number"),
            "title": p.get("title"),
            "status": p.get("status", "not_started"),
            "progress_percentage": p.get("progress_percentage", 0),
            "topics_count": len(p.get("topics", [])),
        } for p in phases],
    }


# ──────────────────────────────────────────────
#  OPTIMIZED RESUME
# ──────────────────────────────────────────────

@router.post("/resume/optimize", response_model=OptimizedResumeResponse)
@limiter.limit("3/minute")
def optimize_resume(
    request: Request,
    resume_analysis_id: int,
    job_description_id: int = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    analysis = db.query(ResumeAnalysis).filter(
        ResumeAnalysis.id == resume_analysis_id,
        ResumeAnalysis.user_id == current_user.id,
    ).first()
    if not analysis:
        raise HTTPException(status_code=404, detail="Resume analysis not found")

    ats_suggestions = json.loads(analysis.ats_suggestions) if analysis.ats_suggestions else []

    jd = None
    if job_description_id:
        jd = db.query(JobDescription).filter(
            JobDescription.id == job_description_id,
            JobDescription.user_id == current_user.id,
        ).first()
        if not jd:
            raise HTTPException(status_code=404, detail="Job description not found")

    resume = db.query(Resume).filter(Resume.id == analysis.resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Original resume not found")

    optimized_data = CareerService.optimize_resume(
        resume.extracted_text,
        ats_suggestions,
        jd.raw_text if jd else None,
    )

    original_ext = os.path.splitext(resume.filename or "")[1].lower() if resume.filename else ""
    resume_format = "docx" if original_ext == ".docx" else "pdf"

    optimized = OptimizedResume(
        user_id=current_user.id,
        resume_analysis_id=analysis.id,
        optimized_text=optimized_data.get("optimized_text", ""),
        improvements=json.dumps(optimized_data.get("improvements", [])),
        professional_summary=optimized_data.get("professional_summary"),
        optimized_skills=json.dumps(optimized_data.get("optimized_skills", [])),
        optimized_projects=json.dumps(optimized_data.get("optimized_projects", [])),
        optimized_keywords=json.dumps(optimized_data.get("optimized_keywords", [])),
        optimized_experience=json.dumps(optimized_data.get("optimized_experience", [])),
        format=resume_format,
    )
    db.add(optimized)
    db.commit()
    db.refresh(optimized)

    # Generate format-preserved file matching the original upload format
    try:
        optimized_text = optimized_data.get("optimized_text", resume.extracted_text)
        analysis_dict = CareerService.calculate_comprehensive_ats_score(
            resume.extracted_text, jd.raw_text if jd else None
        )

        if original_ext == ".docx" and resume.file_path and os.path.exists(resume.file_path):
            docx_bytes = CareerService.optimize_docx_format_preserving(
                resume.file_path, resume.extracted_text, analysis_dict, jd.raw_text if jd else None
            )
            docx_dir = os.path.join("uploads", "optimized")
            os.makedirs(docx_dir, exist_ok=True)
            docx_path = os.path.join(docx_dir, f"optimized_{optimized.id}.docx")
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)
            optimized.file_path = docx_path
            optimized.format = "docx"
        else:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            import io as _io

            output = _io.BytesIO()
            pdf_doc = SimpleDocTemplate(output, pagesize=letter,
                rightMargin=0.5*inch, leftMargin=0.5*inch,
                topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            story = []
            for line in optimized_text.split("\n"):
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                if not safe.strip():
                    story.append(Spacer(1, 4))
                elif safe.strip().isupper() or (len(safe.strip()) < 60 and safe.strip().replace(" ", "").isupper()):
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(safe, styles["Heading2"]))
                else:
                    story.append(Paragraph(safe, styles["Normal"]))
            pdf_doc.build(story)
            output.seek(0)
            pdf_dir = os.path.join("uploads", "optimized")
            os.makedirs(pdf_dir, exist_ok=True)
            pdf_path = os.path.join(pdf_dir, f"optimized_{optimized.id}.pdf")
            with open(pdf_path, "wb") as f:
                f.write(output.getvalue())
            optimized.file_path = pdf_path
            optimized.format = "pdf"

        db.commit()
        db.refresh(optimized)
    except Exception:
        pass

    # Trigger readiness recalculation after optimization
    try:
        from app.services.automation_service import AutomationService
        automation = AutomationService(db)
        automation._update_readiness(current_user.id, trigger_event="resume_optimization")
    except Exception:
        pass

    return optimized


@router.get("/optimized-resume/{opt_id}", response_model=OptimizedResumeResponse)
def get_optimized_resume(
    opt_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    optimized = db.query(OptimizedResume).filter(
        OptimizedResume.id == opt_id,
        OptimizedResume.user_id == current_user.id,
    ).first()
    if not optimized:
        raise HTTPException(status_code=404, detail="Optimized resume not found")
    return optimized


# ──────────────────────────────────────────────
#  CAREER READINESS
# ──────────────────────────────────────────────

@router.get("/readiness", response_model=CareerReadinessResponse)
def get_career_readiness(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    latest_analysis = (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id)
        .order_by(ResumeAnalysis.created_at.desc())
        .first()
    )

    latest_gap = (
        db.query(SkillGapAnalysis)
        .filter(SkillGapAnalysis.user_id == current_user.id)
        .order_by(SkillGapAnalysis.created_at.desc())
        .first()
    )

    latest_interview = (
        db.query(InterviewSession)
        .filter(
            InterviewSession.user_id == current_user.id,
            InterviewSession.status == "completed",
        )
        .order_by(InterviewSession.ended_at.desc())
        .first()
    )

    latest_coding = (
        db.query(CodingSession)
        .filter(
            CodingSession.user_id == current_user.id,
            CodingSession.status == "submitted",
        )
        .order_by(CodingSession.ended_at.desc())
        .first()
    )

    scores = CareerService.calculate_career_readiness(
        resume_match=latest_analysis.resume_match_score if latest_analysis else None,
        ats_score=latest_analysis.ats_score if latest_analysis else None,
        skill_gap=latest_gap.match_percentage if latest_gap else None,
        interview_score=latest_interview.score if latest_interview else None,
        coding_score=latest_coding.coding_score if latest_coding else None,
    )

    existing = (
        db.query(CareerReadiness)
        .filter(CareerReadiness.user_id == current_user.id)
        .order_by(CareerReadiness.created_at.desc())
        .first()
    )

    if existing:
        existing.resume_analysis_id = latest_analysis.id if latest_analysis else existing.resume_analysis_id
        existing.skill_gap_id = latest_gap.id if latest_gap else existing.skill_gap_id
        existing.resume_match_score = scores.get("resume_match_score")
        existing.ats_score = scores.get("ats_score")
        existing.interview_score = scores.get("interview_score")
        existing.coding_score = scores.get("coding_score")
        existing.skill_gap_score = scores.get("skill_gap_score")
        existing.overall_score = scores.get("overall_score")
        existing.recommendations = json.dumps(scores.get("recommendations", []))
        existing.ai_suggestions = json.dumps(scores.get("ai_suggestions", []))
        db.commit()
        db.refresh(existing)
        return existing

    readiness = CareerReadiness(
        user_id=current_user.id,
        resume_analysis_id=latest_analysis.id if latest_analysis else None,
        skill_gap_id=latest_gap.id if latest_gap else None,
        resume_match_score=scores.get("resume_match_score"),
        ats_score=scores.get("ats_score"),
        interview_score=scores.get("interview_score"),
        coding_score=scores.get("coding_score"),
        skill_gap_score=scores.get("skill_gap_score"),
        overall_score=scores.get("overall_score"),
        recommendations=json.dumps(scores.get("recommendations", [])),
        ai_suggestions=json.dumps(scores.get("ai_suggestions", [])),
    )
    db.add(readiness)
    db.commit()
    db.refresh(readiness)
    return readiness


@router.get("/readiness/history", response_model=List[CareerReadinessHistoryResponse])
@limiter.limit("30/minute")
def get_readiness_history(
    request: Request,
    limit: int = Query(30, ge=1, le=200),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get career readiness score history for trend analysis."""
    from app.models.career import CareerReadinessHistory
    history = db.query(CareerReadinessHistory).filter(
        CareerReadinessHistory.user_id == current_user.id,
    ).order_by(CareerReadinessHistory.created_at.desc()).limit(limit).all()
    return history


@router.get("/readiness/breakdown")
@limiter.limit("30/minute")
def get_readiness_breakdown(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get detailed score breakdown with weights and explanations."""
    readiness = db.query(CareerReadiness).filter(
        CareerReadiness.user_id == current_user.id,
    ).order_by(CareerReadiness.created_at.desc()).first()

    if not readiness:
        return {"breakdown": {}, "overall_score": 0}

    breakdown = json.loads(readiness.score_breakdown) if readiness.score_breakdown else {}
    return {
        "overall_score": readiness.overall_score or 0,
        "breakdown": breakdown,
        "weights": {
            "resume_match": 0.20, "ats": 0.10, "skill_gap": 0.20,
            "interview": 0.15, "coding": 0.10, "learning": 0.10,
            "project": 0.05, "consistency": 0.05,
            "role_match": 0.03, "company_match": 0.02,
        },
    }


# ── Role & Company Readiness ──

ROLE_SKILL_REQUIREMENTS = {
    "Backend Developer": ["python", "java", "sql", "rest_api", "docker", "kubernetes", "aws", "git", "postgresql", "mongodb", "redis", "system_design"],
    "Frontend Developer": ["javascript", "typescript", "react", "vue", "angular", "html", "css", "webpack", "git", "testing", "responsive_design"],
    "Full Stack Developer": ["javascript", "typescript", "react", "node.js", "python", "sql", "mongodb", "docker", "aws", "git", "rest_api", "system_design"],
    "Machine Learning Engineer": ["python", "tensorflow", "pytorch", "scikit_learn", "pandas", "numpy", "sql", "docker", "aws", "machine_learning", "deep_learning"],
    "Data Scientist": ["python", "r", "sql", "pandas", "numpy", "scikit_learn", "tableau", "machine_learning", "statistics", "deep_learning"],
    "DevOps Engineer": ["linux", "docker", "kubernetes", "aws", "terraform", "jenkins", "git", "ci_cd", "monitoring", "bash"],
    "Cloud Engineer": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "linux", "networking", "security", "ci_cd"],
    "Software Engineer": ["python", "java", "javascript", "sql", "git", "docker", "rest_api", "data_structures", "algorithms", "system_design"],
}

COMPANY_REQUIREMENTS = {
    "Google": {"min_score": 85, "key_skills": ["system_design", "algorithms", "data_structures", "python", "java"], "focus": "Data structures, algorithms, and system design"},
    "Microsoft": {"min_score": 80, "key_skills": ["csharp", "dotnet", "azure", "system_design", "python"], "focus": "Full-stack development and cloud"},
    "Amazon": {"min_score": 80, "key_skills": ["java", "python", "aws", "system_design", "dsa"], "focus": "Leadership principles and scalable systems"},
    "Meta": {"min_score": 82, "key_skills": ["python", "react", "system_design", "php", "graphql"], "focus": "Large-scale distributed systems"},
    "TCS": {"min_score": 50, "key_skills": ["java", "python", "sql", "spring"], "focus": "Core programming and databases"},
    "Infosys": {"min_score": 50, "key_skills": ["java", "python", "sql", "react"], "focus": "Full-stack development basics"},
    "Accenture": {"min_score": 55, "key_skills": ["python", "java", "sql", "cloud"], "focus": "Consulting and cloud migration"},
    "Wipro": {"min_score": 50, "key_skills": ["java", "python", "sql", "linux"], "focus": "Enterprise application development"},
    "IBM": {"min_score": 60, "key_skills": ["python", "java", "watson", "cloud", "docker"], "focus": "AI/ML and enterprise solutions"},
}


@router.get("/readiness/roles")
@limiter.limit("10/minute")
def get_role_readiness(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Calculate readiness percentage for each target role."""
    latest_gap = db.query(SkillGapAnalysis).filter(
        SkillGapAnalysis.user_id == current_user.id,
    ).order_by(SkillGapAnalysis.created_at.desc()).first()

    latest_analysis = db.query(ResumeAnalysis).filter(
        ResumeAnalysis.user_id == current_user.id,
    ).order_by(ResumeAnalysis.created_at.desc()).first()

    resume_skills = set()
    if latest_analysis and latest_analysis.detected_skills:
        try:
            resume_skills = set(s.lower().strip() for s in json.loads(latest_analysis.detected_skills))
        except (json.JSONDecodeError, TypeError):
            pass

    matched = set()
    if latest_gap and latest_gap.matched_skills:
        try:
            matched = set(s.lower().strip() for s in json.loads(latest_gap.matched_skills))
        except (json.JSONDecodeError, TypeError):
            pass

    all_user_skills = resume_skills | matched

    results = []
    for role, requirements in ROLE_SKILL_REQUIREMENTS.items():
        req_set = set(r.lower().strip() for r in requirements)
        matched_count = len(all_user_skills & req_set)
        readiness_score = round((matched_count / max(1, len(req_set))) * 100, 1)

        matched_list = list(all_user_skills & req_set)
        missing_list = list(req_set - all_user_skills)

        recommendations = []
        if missing_list:
            recommendations.append(f"Learn {', '.join(missing_list[:3])} to improve {role} readiness")
        if readiness_score < 50:
            recommendations.append(f"Focus on core {role} skills before applying")
        elif readiness_score < 75:
            recommendations.append(f"Good foundation — strengthen {', '.join(missing_list[:2])} to reach 75%+")

        results.append({
            "role": role,
            "readiness_score": readiness_score,
            "matched_skills": matched_list,
            "missing_skills": missing_list,
            "recommendations": recommendations,
        })

    results.sort(key=lambda x: x["readiness_score"], reverse=True)
    return results


@router.get("/readiness/companies")
@limiter.limit("10/minute")
def get_company_readiness(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Calculate readiness percentage for each target company."""
    latest_gap = db.query(SkillGapAnalysis).filter(
        SkillGapAnalysis.user_id == current_user.id,
    ).order_by(SkillGapAnalysis.created_at.desc()).first()

    latest_analysis = db.query(ResumeAnalysis).filter(
        ResumeAnalysis.user_id == current_user.id,
    ).order_by(ResumeAnalysis.created_at.desc()).first()

    resume_skills = set()
    if latest_analysis and latest_analysis.detected_skills:
        try:
            resume_skills = set(s.lower().strip() for s in json.loads(latest_analysis.detected_skills))
        except (json.JSONDecodeError, TypeError):
            pass

    matched = set()
    if latest_gap and latest_gap.matched_skills:
        try:
            matched = set(s.lower().strip() for s in json.loads(latest_gap.matched_skills))
        except (json.JSONDecodeError, TypeError):
            pass

    all_user_skills = resume_skills | matched

    # Also consider interview and coding scores
    latest_interview = db.query(InterviewSession).filter(
        InterviewSession.user_id == current_user.id,
        InterviewSession.status == "completed",
    ).order_by(InterviewSession.ended_at.desc()).first()

    latest_coding = db.query(CodingSession).filter(
        CodingSession.user_id == current_user.id,
        CodingSession.status == "submitted",
    ).order_by(CodingSession.ended_at.desc()).first()

    interview_score = latest_interview.score if latest_interview else 0
    coding_score = latest_coding.coding_score if latest_coding else 0

    results = []
    for company, reqs in COMPANY_REQUIREMENTS.items():
        key_skills = set(s.lower().strip() for s in reqs["key_skills"])
        matched_count = len(all_user_skills & key_skills)
        skill_score = (matched_count / max(1, len(key_skills))) * 100

        # Weighted score: 50% skills + 25% interview + 25% coding
        readiness_score = round(
            skill_score * 0.50 + interview_score * 0.25 + coding_score * 0.25, 1
        )
        readiness_score = min(100, readiness_score)

        matched_list = list(all_user_skills & key_skills)
        missing_list = list(key_skills - all_user_skills)

        recommendations = []
        if readiness_score < reqs["min_score"]:
            gap = reqs["min_score"] - readiness_score
            recommendations.append(f"Need {round(gap, 1)}% more to meet {company}'s bar")
        if missing_list:
            recommendations.append(f"Strengthen: {', '.join(missing_list[:3])}")
        recommendations.append(f"Focus area: {reqs['focus']}")

        results.append({
            "company": company,
            "readiness_score": readiness_score,
            "matched_skills": matched_list,
            "missing_skills": missing_list,
            "recommendations": recommendations,
        })

    results.sort(key=lambda x: x["readiness_score"], reverse=True)
    return results


@router.post("/readiness/recalculate")
@limiter.limit("10/minute")
def recalculate_readiness(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Manually trigger career readiness recalculation."""
    from app.services.automation_service import AutomationService
    automation = AutomationService(db)
    automation._update_readiness(current_user.id, trigger_event="manual")

    readiness = db.query(CareerReadiness).filter(
        CareerReadiness.user_id == current_user.id,
    ).order_by(CareerReadiness.created_at.desc()).first()

    return readiness or {"overall_score": 0, "message": "No readiness data found"}


# ──────────────────────────────────────────────
#  DASHBOARD
# ──────────────────────────────────────────────

@router.get("/dashboard", response_model=CareerDashboardResponse)
def get_career_dashboard(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    latest_analysis = (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id)
        .order_by(ResumeAnalysis.created_at.desc())
        .first()
    )

    latest_gap = (
        db.query(SkillGapAnalysis)
        .filter(SkillGapAnalysis.user_id == current_user.id)
        .order_by(SkillGapAnalysis.created_at.desc())
        .first()
    )

    latest_readiness = (
        db.query(CareerReadiness)
        .filter(CareerReadiness.user_id == current_user.id)
        .order_by(CareerReadiness.created_at.desc())
        .first()
    )

    recent_analyses = (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id)
        .order_by(ResumeAnalysis.created_at.desc())
        .limit(5)
        .all()
    )

    roadmap = (
        db.query(LearningRoadmap)
        .filter(
            LearningRoadmap.user_id == current_user.id,
            LearningRoadmap.status == "active",
        )
        .order_by(LearningRoadmap.created_at.desc())
        .first()
    )

    latest_interview = (
        db.query(InterviewSession)
        .filter(
            InterviewSession.user_id == current_user.id,
            InterviewSession.status == "completed",
        )
        .order_by(InterviewSession.ended_at.desc())
        .first()
    )

    latest_coding = (
        db.query(CodingSession)
        .filter(
            CodingSession.user_id == current_user.id,
            CodingSession.status == "submitted",
        )
        .order_by(CodingSession.ended_at.desc())
        .first()
    )

    interview_analyses = (
        db.query(InterviewSession)
        .filter(
            InterviewSession.user_id == current_user.id,
            InterviewSession.status == "completed",
        )
        .order_by(InterviewSession.ended_at.desc())
        .limit(10)
        .all()
    )

    coding_analyses = (
        db.query(CodingSession)
        .filter(
            CodingSession.user_id == current_user.id,
            CodingSession.status == "submitted",
        )
        .order_by(CodingSession.ended_at.desc())
        .limit(10)
        .all()
    )

    resume_match = latest_analysis.resume_match_score if latest_analysis else None
    ats_score = latest_analysis.ats_score if latest_analysis else None
    missing_skills = json.loads(latest_gap.missing_skills) if latest_gap and latest_gap.missing_skills else []
    recommendations = json.loads(latest_readiness.recommendations) if latest_readiness and latest_readiness.recommendations else []
    ai_suggestions = json.loads(latest_readiness.ai_suggestions) if latest_readiness and latest_readiness.ai_suggestions else []

    resume_match_trend = []
    all_analyses = (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id, ResumeAnalysis.resume_match_score.isnot(None))
        .order_by(ResumeAnalysis.created_at.asc())
        .all()
    )
    for a in all_analyses:
        resume_match_trend.append({
            "date": a.created_at.isoformat() if a.created_at else None,
            "score": a.resume_match_score,
        })

    interview_trend = []
    for s in interview_analyses:
        interview_trend.append({
            "date": s.ended_at.isoformat() if s.ended_at else None,
            "score": s.score,
        })

    coding_trend = []
    for c in coding_analyses:
        coding_trend.append({
            "date": c.ended_at.isoformat() if c.ended_at else None,
            "score": c.coding_score,
        })

    return CareerDashboardResponse(
        resume_match_score=resume_match,
        ats_score=ats_score,
        career_readiness=latest_readiness.overall_score if latest_readiness else None,
        interview_readiness=latest_interview.score if latest_interview else None,
        coding_readiness=latest_coding.coding_score if latest_coding else None,
        missing_skills=missing_skills,
        recent_analyses=[ResumeAnalysisResponse.model_validate(a) for a in recent_analyses],
        skill_gap=SkillGapResponse.model_validate(latest_gap) if latest_gap else None,
        roadmap=LearningRoadmapResponse.model_validate(roadmap) if roadmap else None,
        recommendations=recommendations,
        ai_suggestions=ai_suggestions,
        resume_match_trend=resume_match_trend,
        interview_trend=interview_trend,
        coding_trend=coding_trend,
    )


# ──────────────────────────────────────────────
#  AI SUGGESTIONS
# ──────────────────────────────────────────────

@router.get("/suggestions")
def get_ai_suggestions(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    latest_analysis = (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id)
        .order_by(ResumeAnalysis.created_at.desc())
        .first()
    )

    latest_gap = (
        db.query(SkillGapAnalysis)
        .filter(SkillGapAnalysis.user_id == current_user.id)
        .order_by(SkillGapAnalysis.created_at.desc())
        .first()
    )

    resume_skills = json.loads(latest_analysis.detected_skills) if latest_analysis and latest_analysis.detected_skills else []
    missing = json.loads(latest_gap.missing_skills) if latest_gap and latest_gap.missing_skills else []

    suggestions = CareerService.get_ai_suggestions(
        resume_skills=resume_skills,
        missing_skills=missing,
        ats_score=latest_analysis.ats_score if latest_analysis else None,
        match_score=latest_analysis.resume_match_score if latest_analysis else None,
    )

    return suggestions


@router.delete("/job-description/{jd_id}")
def delete_job_description(
    jd_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    jd = db.query(JobDescription).filter(
        JobDescription.id == jd_id,
        JobDescription.user_id == current_user.id,
    ).first()
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found")

    sg_ids = [sg.id for sg in db.query(SkillGapAnalysis.id).filter(
        SkillGapAnalysis.job_description_id == jd_id
    ).all()]
    if sg_ids:
        db.query(LearningRoadmap).filter(
            LearningRoadmap.skill_gap_id.in_(sg_ids)
        ).delete(synchronize_session=False)
        db.query(SkillGapAnalysis).filter(
            SkillGapAnalysis.id.in_(sg_ids)
        ).delete(synchronize_session=False)

    ra_ids = [ra.id for ra in db.query(ResumeAnalysis.id).filter(
        ResumeAnalysis.job_description_id == jd_id
    ).all()]
    if ra_ids:
        db.query(SkillGapAnalysis).filter(
            SkillGapAnalysis.resume_analysis_id.in_(ra_ids)
        ).delete(synchronize_session=False)
        db.query(OptimizedResume).filter(
            OptimizedResume.resume_analysis_id.in_(ra_ids)
        ).delete(synchronize_session=False)
        db.query(CareerReadiness).filter(
            CareerReadiness.resume_analysis_id.in_(ra_ids)
        ).delete(synchronize_session=False)
        db.query(CareerReadinessHistory).filter(
            CareerReadinessHistory.user_id == current_user.id
        ).delete(synchronize_session=False)
        db.query(ResumeAnalysis).filter(
            ResumeAnalysis.id.in_(ra_ids)
        ).delete(synchronize_session=False)

    if jd.file_path and os.path.exists(jd.file_path):
        os.remove(jd.file_path)

    db.delete(jd)
    db.commit()
    return {"message": "Job description and all related records deleted."}


@router.delete("/resume-analysis/{analysis_id}")
def delete_resume_analysis(
    analysis_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    analysis = db.query(ResumeAnalysis).filter(
        ResumeAnalysis.id == analysis_id,
        ResumeAnalysis.user_id == current_user.id,
    ).first()
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    # Cascade: delete related records
    db.query(OptimizedResume).filter(
        OptimizedResume.resume_analysis_id == analysis_id
    ).delete(synchronize_session=False)
    db.query(CareerReadiness).filter(
        CareerReadiness.resume_analysis_id == analysis_id
    ).delete(synchronize_session=False)
    db.query(CareerReadinessHistory).filter(
        CareerReadinessHistory.user_id == current_user.id
    ).delete(synchronize_session=False)

    db.delete(analysis)
    db.commit()
    return {"message": "Analysis deleted"}


@router.delete("/roadmap/{roadmap_id}")
def delete_roadmap(
    roadmap_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    roadmap = db.query(LearningRoadmap).filter(
        LearningRoadmap.id == roadmap_id,
        LearningRoadmap.user_id == current_user.id,
    ).first()
    if not roadmap:
        raise HTTPException(status_code=404, detail="Roadmap not found")
    db.delete(roadmap)
    db.commit()
    return {"message": "Roadmap deleted"}


@router.delete("/optimized-resume/{opt_id}")
def delete_optimized_resume(
    opt_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    optimized = db.query(OptimizedResume).filter(
        OptimizedResume.id == opt_id,
        OptimizedResume.user_id == current_user.id,
    ).first()
    if not optimized:
        raise HTTPException(status_code=404, detail="Optimized resume not found")
    db.delete(optimized)
    db.commit()
    return {"message": "Optimized resume deleted"}


class RoadmapProgressRequest(BaseModel):
    completed_topics: list[str] = []


@router.patch("/roadmap/{roadmap_id}/progress")
@limiter.limit("30/minute")
def update_roadmap_progress(
    request: Request,
    roadmap_id: int,
    completed_topics: list[str] = Body(default=[], embed=True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    roadmap = db.query(LearningRoadmap).filter(
        LearningRoadmap.id == roadmap_id,
        LearningRoadmap.user_id == current_user.id,
    ).first()
    if not roadmap:
        raise HTTPException(status_code=404, detail="Roadmap not found")
    
    # Use LearningProgressTracker for proper tracking (phase auto-advance, skill analytics, readiness)
    from app.models.intelligence import LearningProgress
    from app.services.learning_progress_tracker import LearningProgressTracker
    tracker = LearningProgressTracker(db)

    # Get existing progress records to avoid duplicates
    existing_progress = {
        p.topic_name for p in db.query(LearningProgress.topic_name).filter(
            LearningProgress.user_id == current_user.id,
            LearningProgress.roadmap_id == roadmap_id,
        ).all()
    }

    # Mark newly completed topics via tracker (updates skill analytics + readiness)
    for topic_name in completed_topics:
        if topic_name not in existing_progress:
            try:
                tracker.start_topic(current_user.id, roadmap_id, topic_name)
                tracker.complete_topic(current_user.id, roadmap_id, topic_name)
            except Exception:
                pass

    # Mark topics that were unchecked (removed from completed list)
    for topic_name in existing_progress:
        if topic_name not in completed_topics:
            try:
                progress = db.query(LearningProgress).filter(
                    LearningProgress.user_id == current_user.id,
                    LearningProgress.roadmap_id == roadmap_id,
                    LearningProgress.topic_name == topic_name,
                ).first()
                if progress:
                    progress.status = "not_started"
                    progress.progress_percentage = 0.0
                    progress.completed_at = None
                    db.commit()
            except Exception:
                pass

    # Set completed_topics AFTER tracker calls (tracker._update_roadmap_progress overwrites it)
    roadmap.completed_topics = json.dumps(completed_topics)

    # Count all individual topics across all phases for accurate percentage
    phases = json.loads(roadmap.phases) if roadmap.phases else []
    
    total_topics = 0
    for phase in phases:
        for t in phase.get("topics", []):
            total_topics += 1
    
    done = len(completed_topics)
    roadmap.progress_percentage = round((done / max(total_topics, 1)) * 100, 1)

    # Auto-advance current_phase_index
    current_idx = roadmap.current_phase_index or 0
    if phases and current_idx < len(phases):
        current_phase = phases[current_idx]
        phase_topic_names = [t.get("name", t) if isinstance(t, dict) else t for t in current_phase.get("topics", [])]
        if phase_topic_names and all(t in completed_topics for t in phase_topic_names):
            if current_idx < len(phases) - 1:
                roadmap.current_phase_index = current_idx + 1
    
    if done == total_topics and total_topics > 0:
        roadmap.status = "completed"
    
    db.commit()
    db.refresh(roadmap)

    # Return roadmap with parsed JSON fields
    return {
        "id": roadmap.id,
        "progress_percentage": roadmap.progress_percentage,
        "current_phase_index": roadmap.current_phase_index,
        "status": roadmap.status,
        "completed_topics": completed_topics,
    }


# ──────────────────────────────────────────────
#  ROADMAP HISTORY & CURRENT
# ──────────────────────────────────────────────


# ──────────────────────────────────────────────
#  INTELLIGENCE LAYER — Skill Gap Engine
# ──────────────────────────────────────────────

@router.post("/skill-gap/analyze-full")
def analyze_skill_gap_full(
    resume_analysis_id: int = None,
    job_description_id: int = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Full intelligence pipeline: resume + JD + interview + coding → unified skill assessment."""
    from app.services.skill_gap_engine import SkillGapEngine

    engine = SkillGapEngine(db)
    result = engine.analyze(
        user_id=current_user.id,
        resume_analysis_id=resume_analysis_id,
        job_description_id=job_description_id,
    )

    analysis = engine.save_to_db(
        current_user.id, result,
        resume_analysis_id=resume_analysis_id,
        job_description_id=job_description_id,
    )
    engine.update_skill_analytics(current_user.id, result)
    engine.generate_career_recommendations(current_user.id, result)

    try:
        from app.services.automation_service import AutomationService
        automation = AutomationService(db)
        automation._update_readiness(current_user.id, trigger_event="skill_gap")
    except Exception:
        pass

    result["skill_gap_id"] = analysis.id
    return result


@router.post("/skill-gap/analyze")
def analyze_skill_gap(
    resume_analysis_id: int = None,
    job_description_id: int = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Alias for analyze-full — more discoverable. Full intelligence pipeline: resume + JD + interview + coding → unified skill assessment."""
    from app.services.skill_gap_engine import SkillGapEngine

    engine = SkillGapEngine(db)
    result = engine.analyze(
        user_id=current_user.id,
        resume_analysis_id=resume_analysis_id,
        job_description_id=job_description_id,
    )

    analysis = engine.save_to_db(
        current_user.id, result,
        resume_analysis_id=resume_analysis_id,
        job_description_id=job_description_id,
    )
    engine.update_skill_analytics(current_user.id, result)
    engine.generate_career_recommendations(current_user.id, result)

    try:
        from app.services.automation_service import AutomationService
        automation = AutomationService(db)
        automation._update_readiness(current_user.id, trigger_event="skill_gap")
    except Exception:
        pass

    result["skill_gap_id"] = analysis.id
    return result


@router.get("/skill-gap-trend")
def get_skill_gap_trend(
    limit: int = Query(10, ge=1, le=50),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get skill gap trend over time — compare latest analyses."""
    analyses = (
        db.query(SkillGapAnalysis)
        .filter(SkillGapAnalysis.user_id == current_user.id)
        .order_by(SkillGapAnalysis.created_at.desc())
        .limit(limit)
        .all()
    )

    if not analyses:
        return {"trend": [], "summary": None}

    trend = []
    prev_match = None
    for a in analyses:
        matched = json.loads(a.matched_skills) if a.matched_skills else []
        missing = json.loads(a.missing_skills) if a.missing_skills else []
        additional = json.loads(a.additional_skills) if a.additional_skills else []
        priority = json.loads(a.priority_skills) if a.priority_skills else []

        delta = None
        if prev_match is not None and a.match_percentage is not None:
            delta = round(a.match_percentage - prev_match, 1)

        trend.append({
            "id": a.id,
            "date": a.created_at.isoformat() if a.created_at else None,
            "match_percentage": a.match_percentage,
            "delta": delta,
            "matched_count": len(matched),
            "missing_count": len(missing),
            "additional_count": len(additional),
            "priority_count": len(priority),
            "resume_analysis_id": a.resume_analysis_id,
            "job_description_id": a.job_description_id,
        })

        if a.match_percentage is not None:
            prev_match = a.match_percentage

    # Summary stats
    matches = [t["match_percentage"] for t in trend if t["match_percentage"] is not None]
    deltas = [t["delta"] for t in trend if t["delta"] is not None]

    summary = {
        "total_analyses": len(trend),
        "latest_match": matches[0] if matches else None,
        "best_match": max(matches) if matches else None,
        "worst_match": min(matches) if matches else None,
        "avg_match": round(sum(matches) / len(matches), 1) if matches else None,
        "improvement": round(deltas[0], 1) if deltas else None,
        "avg_improvement": round(sum(deltas) / len(deltas), 1) if deltas else None,
        "trend_direction": "improving" if deltas and deltas[0] > 0 else "declining" if deltas and deltas[0] < 0 else "stable",
    }

    return {"trend": trend, "summary": summary}


@router.get("/learning-progress")
@limiter.limit("30/minute")
def get_learning_progress(
    request: Request,
    roadmap_id: int = None,
    page: int = Query(1, ge=1),
    per_page: int = Query(50, ge=1, le=200),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get learning progress for current user with pagination."""
    from app.services.learning_progress_tracker import LearningProgressTracker

    tracker = LearningProgressTracker(db)
    all_progress = tracker.get_progress(current_user.id, roadmap_id)
    stats = tracker.get_stats(current_user.id)

    total = len(all_progress)
    start = (page - 1) * per_page
    end = start + per_page
    paginated = all_progress[start:end]

    return {
        "progress": paginated,
        "stats": stats,
        "pagination": {
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page,
        },
    }


@router.post("/learning-progress/start")
@limiter.limit("30/minute")
def start_learning_topic(
    request: Request,
    roadmap_id: int,
    topic_name: str,
    skill_name: str = None,
    estimated_hours: float = None,
    difficulty: str = None,
    priority: str = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Mark a learning topic as in-progress."""
    from app.services.learning_progress_tracker import LearningProgressTracker

    tracker = LearningProgressTracker(db)
    progress = tracker.start_topic(
        current_user.id, roadmap_id, topic_name,
        skill_name, estimated_hours, difficulty, priority,
    )
    return {"id": progress.id, "status": progress.status, "message": f"Started: {topic_name}"}


@router.post("/learning-progress/complete")
@limiter.limit("30/minute")
def complete_learning_topic(
    request: Request,
    roadmap_id: int,
    topic_name: str,
    actual_hours: float = None,
    notes: str = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Mark a learning topic as completed and update readiness."""
    from app.services.learning_progress_tracker import LearningProgressTracker

    tracker = LearningProgressTracker(db)
    progress = tracker.complete_topic(
        current_user.id, roadmap_id, topic_name, actual_hours, notes,
    )
    stats = tracker.get_stats(current_user.id)
    return {
        "id": progress.id,
        "status": progress.status,
        "message": f"Completed: {topic_name}",
        "stats": stats,
    }


@router.post("/learning-progress/master")
@limiter.limit("30/minute")
def master_learning_topic(
    request: Request,
    roadmap_id: int,
    topic_name: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Mark a learning topic as mastered."""
    from app.services.learning_progress_tracker import LearningProgressTracker

    tracker = LearningProgressTracker(db)
    progress = tracker.master_topic(current_user.id, roadmap_id, topic_name)
    return {"id": progress.id, "status": progress.status, "message": f"Mastered: {topic_name}"}


@router.put("/learning-progress/update")
@limiter.limit("30/minute")
def update_learning_progress(
    request: Request,
    roadmap_id: int,
    topic_name: str,
    percentage: float,
    actual_hours: float = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Update partial progress on a learning topic."""
    from app.services.learning_progress_tracker import LearningProgressTracker

    tracker = LearningProgressTracker(db)
    progress = tracker.update_progress(
        current_user.id, roadmap_id, topic_name, percentage, actual_hours,
    )
    return {"id": progress.id, "progress_percentage": progress.progress_percentage}


@router.get("/recommendations")
def get_career_recommendations(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get prioritized career recommendations."""
    from app.models.intelligence import CareerRecommendation

    recs = db.query(CareerRecommendation).filter(
        CareerRecommendation.user_id == current_user.id,
        CareerRecommendation.is_dismissed == False,
    ).order_by(
        CareerRecommendation.priority.desc(),
        CareerRecommendation.created_at.desc(),
    ).all()
    return recs


@router.put("/recommendations/{rec_id}/dismiss")
def dismiss_recommendation(
    rec_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Dismiss a career recommendation."""
    from app.models.intelligence import CareerRecommendation

    rec = db.query(CareerRecommendation).filter(
        CareerRecommendation.id == rec_id,
        CareerRecommendation.user_id == current_user.id,
    ).first()
    if not rec:
        raise HTTPException(status_code=404, detail="Recommendation not found")
    rec.is_dismissed = True
    db.commit()
    return {"message": "Recommendation dismissed"}


@router.put("/recommendations/{rec_id}/complete")
def complete_recommendation(
    rec_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Mark a career recommendation as completed."""
    from app.models.intelligence import CareerRecommendation

    rec = db.query(CareerRecommendation).filter(
        CareerRecommendation.id == rec_id,
        CareerRecommendation.user_id == current_user.id,
    ).first()
    if not rec:
        raise HTTPException(status_code=404, detail="Recommendation not found")
    rec.is_completed = True
    db.commit()
    return {"message": "Recommendation completed"}


@router.get("/analytics/skill-heatmap")
def get_skill_heatmap(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get skill analytics heatmap for visualization."""
    from app.models.intelligence import SkillAnalytics

    skills = db.query(SkillAnalytics).filter(
        SkillAnalytics.user_id == current_user.id,
    ).all()

    categories = {}
    total_proficiency = 0
    for skill in skills:
        cat = skill.category or "other"
        if cat not in categories:
            categories[cat] = {"count": 0, "avg_proficiency": 0, "skills": []}
        categories[cat]["count"] += 1
        categories[cat]["skills"].append({
            "name": skill.skill_name,
            "proficiency": skill.proficiency_level,
            "trend": skill.trend,
        })
        total_proficiency += skill.proficiency_level

    # Calculate averages
    for cat in categories:
        profs = [s["proficiency"] for s in categories[cat]["skills"]]
        categories[cat]["avg_proficiency"] = round(sum(profs) / max(len(profs), 1), 1)

    overall = round(total_proficiency / max(len(skills), 1), 1)

    return {
        "skills": skills,
        "categories": categories,
        "overall_proficiency": overall,
    }


@router.get("/analytics/performance")
def get_performance_metrics(
    metric_type: str = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get performance metrics from interviews and coding rounds."""
    from app.models.intelligence import PerformanceMetrics

    query = db.query(PerformanceMetrics).filter(
        PerformanceMetrics.user_id == current_user.id,
    )
    if metric_type:
        query = query.filter(PerformanceMetrics.metric_type == metric_type)

    metrics = query.order_by(PerformanceMetrics.created_at.desc()).limit(50).all()
    return metrics


@router.get("/skill-dependencies/{skill_name}")
def get_skill_dependencies(
    skill_name: str,
    current_user: User = Depends(get_current_user),
):
    """Get prerequisite chain for a skill."""
    from app.services.skill_dependency_mapper import SkillDependencyMapper

    mapper = SkillDependencyMapper()
    prereqs = mapper.get_all_prerequisites(skill_name)
    dependents = mapper.get_dependents(skill_name)
    learning_path = mapper.get_learning_path([skill_name])

    return {
        "skill": skill_name,
        "prerequisites": prereqs,
        "dependents": dependents,
        "learning_path": learning_path,
    }


@router.get("/report/pdf")
def generate_career_report_pdf(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    latest_analysis = (
        db.query(ResumeAnalysis)
        .filter(ResumeAnalysis.user_id == current_user.id)
        .order_by(ResumeAnalysis.created_at.desc())
        .first()
    )

    latest_gap = (
        db.query(SkillGapAnalysis)
        .filter(SkillGapAnalysis.user_id == current_user.id)
        .order_by(SkillGapAnalysis.created_at.desc())
        .first()
    )

    latest_readiness = (
        db.query(CareerReadiness)
        .filter(CareerReadiness.user_id == current_user.id)
        .order_by(CareerReadiness.created_at.desc())
        .first()
    )

    roadmap = (
        db.query(LearningRoadmap)
        .filter(
            LearningRoadmap.user_id == current_user.id,
            LearningRoadmap.status == "active",
        )
        .order_by(LearningRoadmap.created_at.desc())
        .first()
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Career Readiness Report", styles['Title']))
    elements.append(Spacer(1, 12))

    data = [
        ["Platform", "AI Career Preparation Platform"],
        ["Candidate", current_user.name],
        ["Date", "Generated on request"],
    ]
    t = Table(data, colWidths=[150, 300])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 24))

    elements.append(Paragraph("Scores Overview", styles['Heading2']))
    score_data = [["Metric", "Score"]]
    if latest_analysis:
        if latest_analysis.resume_match_score is not None:
            score_data.append(["Resume Match", f"{latest_analysis.resume_match_score:.1f}%"])
        if latest_analysis.ats_score is not None:
            score_data.append(["ATS Score", f"{latest_analysis.ats_score:.1f}%"])
    if latest_gap and latest_gap.match_percentage is not None:
        score_data.append(["Skill Gap Match", f"{latest_gap.match_percentage:.1f}%"])
    if latest_readiness:
        if latest_readiness.overall_score is not None:
            score_data.append(["Overall Career Readiness", f"{latest_readiness.overall_score:.1f}%"])
        if latest_readiness.interview_score is not None:
            score_data.append(["Interview Score", f"{latest_readiness.interview_score:.1f}%"])
        if latest_readiness.coding_score is not None:
            score_data.append(["Coding Score", f"{latest_readiness.coding_score:.1f}%"])

    if len(score_data) > 1:
        st = Table(score_data, colWidths=[200, 150])
        st.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.black),
        ]))
        elements.append(st)
    else:
        elements.append(Paragraph("No score data available yet. Complete analyses to generate scores.", styles['Normal']))
    elements.append(Spacer(1, 24))

    if latest_analysis and latest_analysis.summary:
        elements.append(Paragraph("Resume Summary", styles['Heading2']))
        elements.append(Paragraph(latest_analysis.summary, styles['Normal']))
        elements.append(Spacer(1, 12))

    if latest_analysis and latest_analysis.detected_skills:
        skills = json.loads(latest_analysis.detected_skills)
        elements.append(Paragraph("Detected Skills", styles['Heading2']))
        elements.append(Paragraph(", ".join(skills), styles['Normal']))
        elements.append(Spacer(1, 12))

    if latest_gap:
        elements.append(Paragraph("Skill Gap Analysis", styles['Heading2']))
        if latest_gap.matched_skills:
            matched = json.loads(latest_gap.matched_skills)
            elements.append(Paragraph(f"<b>Matched ({len(matched)}):</b> {', '.join(matched)}", styles['Normal']))
        if latest_gap.missing_skills:
            missing = json.loads(latest_gap.missing_skills)
            elements.append(Paragraph(f"<b>Missing ({len(missing)}):</b> {', '.join(missing)}", styles['Normal']))
        if latest_gap.additional_skills:
            additional = json.loads(latest_gap.additional_skills)
            elements.append(Paragraph(f"<b>Additional ({len(additional)}):</b> {', '.join(additional)}", styles['Normal']))
        elements.append(Spacer(1, 12))

    if latest_analysis and latest_analysis.ats_suggestions:
        suggestions = json.loads(latest_analysis.ats_suggestions)
        elements.append(Paragraph("ATS Suggestions", styles['Heading2']))
        for i, s in enumerate(suggestions, 1):
            elements.append(Paragraph(f"{i}. {s}", styles['Normal']))
        elements.append(Spacer(1, 12))

    if roadmap and roadmap.roadmap_items:
        items = json.loads(roadmap.roadmap_items)
        elements.append(Paragraph("Learning Roadmap", styles['Heading2']))
        elements.append(Paragraph(f"Total Hours: {roadmap.total_hours or 0} | Estimated Weeks: {roadmap.estimated_weeks or 0}", styles['Normal']))
        elements.append(Spacer(1, 6))
        for item in items:
            elements.append(Paragraph(f"<b>{item.get('topic', 'N/A')}</b> ({item.get('hours', 0)}h - {item.get('difficulty', 'N/A')})", styles['Heading3']))
            elements.append(Paragraph(item.get('description', ''), styles['Normal']))
            if item.get('mini_project'):
                elements.append(Paragraph(f"<b>Mini Project:</b> {item['mini_project']}", styles['Normal']))
            elements.append(Spacer(1, 6))

    if latest_readiness and latest_readiness.recommendations:
        recs = json.loads(latest_readiness.recommendations)
        elements.append(Paragraph("Recommendations", styles['Heading2']))
        for i, r in enumerate(recs, 1):
            elements.append(Paragraph(f"{i}. {r}", styles['Normal']))
        elements.append(Spacer(1, 12))

    if latest_readiness and latest_readiness.ai_suggestions:
        ai_suggs = json.loads(latest_readiness.ai_suggestions)
        elements.append(Paragraph("AI Suggestions", styles['Heading2']))
        for i, s in enumerate(ai_suggs, 1):
            elements.append(Paragraph(f"{i}. {s}", styles['Normal']))

    doc.build(elements)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=career_report_{current_user.id}.pdf"},
    )
