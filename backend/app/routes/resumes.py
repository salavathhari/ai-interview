import os
import re
import uuid
import hashlib
from datetime import datetime, timezone
from typing import Optional, List

from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, status, Query
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import func as sql_func

from app.database import get_db
from app.auth.utils import get_current_user
from app.models.user import User
from app.models.resume import Resume, ResumeVersion
from app.schemas.resume import ResumeResponse
from app.services.ai_service import AIService
import pdfplumber
from docx import Document
from pathlib import Path

router = APIRouter(
    prefix="/resumes",
    tags=["resumes"]
)

UPLOAD_DIR = Path("uploads/resumes")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _sanitize_filename(filename: str) -> str:
    name = os.path.basename(filename)
    name = re.sub(r"[^a-zA-Z0-9._-]", "_", name)
    name = re.sub(r"_+", "_", name).strip("_.")
    if not name:
        name = "resume"
    return name


def _detect_encrypted_pdf(file_path: Path) -> bool:
    try:
        with open(file_path, "rb") as f:
            header = f.read(20)
            if b"%PDF" not in header:
                return False
            f.seek(0)
            content = f.read()
            if b"/Encrypt" in content or b"/Standard" in content:
                return True
    except Exception:
        pass
    return False


def _extract_text_from_pdf(file_path: Path) -> str:
    if _detect_encrypted_pdf(file_path):
        raise ValueError(
            "This PDF is encrypted or password-protected. "
            "Please upload an unprotected PDF."
        )
    text = ""
    with pdfplumber.open(file_path) as pdf:
        if len(pdf.pages) == 0:
            raise ValueError("PDF has no pages.")
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    return text


def _extract_text_from_docx(file_path: Path) -> str:
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text += row_text + "\n"
    return text


def _extract_structured_fields(text: str) -> dict:
    fields = {
        "name": None,
        "email": None,
        "phone": None,
        "location": None,
        "linkedin": None,
        "github": None,
        "portfolio": None,
    }

    email_match = re.search(
        r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text
    )
    if email_match:
        fields["email"] = email_match.group(0)

    phone_match = re.search(
        r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}",
        text[:500],
    )
    if phone_match:
        phone = phone_match.group(0).strip()
        digits = re.sub(r"\D", "", phone)
        if 7 <= len(digits) <= 15:
            fields["phone"] = phone

    linkedin_match = re.search(
        r"linkedin\.com/in/[a-zA-Z0-9_-]+", text, re.IGNORECASE
    )
    if linkedin_match:
        fields["linkedin"] = "https://" + linkedin_match.group(0)

    github_match = re.search(
        r"github\.com/[a-zA-Z0-9_-]+", text, re.IGNORECASE
    )
    if github_match:
        fields["github"] = "https://" + github_match.group(0)

    portfolio_match = re.search(
        r"(?:portfolio|website|blog)\s*[:\-]?\s*(https?://[^\s,]+)",
        text,
        re.IGNORECASE,
    )
    if portfolio_match:
        fields["portfolio"] = portfolio_match.group(1)

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        first_line = lines[0]
        if not re.search(r"[@]", first_line) and len(first_line) < 60:
            if not any(
                kw in first_line.lower()
                for kw in ["resume", "cv", "curriculum", "contact", "phone"]
            ):
                fields["name"] = first_line

    location_patterns = [
        r"(?:location|address|city)\s*[:\-]?\s*([A-Z][a-zA-Z\s]+(?:,\s*[A-Z]{2,}))",
        r"\b([A-Z][a-z]+(?:,\s*[A-Z]{2,}))\b",
    ]
    for pattern in location_patterns:
        loc_match = re.search(pattern, text[:1000])
        if loc_match:
            fields["location"] = loc_match.group(1).strip()
            break

    return fields


def _compute_content_hash(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()


@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No filename provided.",
        )

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX files are allowed.",
        )

    if file.content_type and file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid file type: {file.content_type}",
        )

    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB.",
        )
    if len(contents) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Empty file uploaded.",
        )

    safe_name = _sanitize_filename(file.filename)
    unique_name = f"{current_user.id}_{uuid.uuid4().hex[:12]}_{safe_name}"
    file_path = UPLOAD_DIR / unique_name

    try:
        with file_path.open("wb") as buffer:
            buffer.write(contents)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Could not save file: {str(e)}",
        )

    try:
        text = ""
        if ext == ".pdf":
            text = _extract_text_from_pdf(file_path)
        elif ext == ".docx":
            text = _extract_text_from_docx(file_path)

        if not text.strip():
            raise ValueError("No text could be extracted from the document.")

    except ValueError:
        if file_path.exists():
            file_path.unlink()
        raise
    except Exception as e:
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not extract text from document: {str(e)}",
        )

    structured_fields = _extract_structured_fields(text)

    try:
        from app.ml.prediction.skill_service import SkillService
        ml_skills = SkillService.extract_skills(text)
        detected_skills = ml_skills.get("all_skills", [])
        skills_str = ", ".join(detected_skills) if detected_skills else None
    except Exception:
        detected_skills = AIService.extract_skills(text)
        skills_str = ", ".join(detected_skills) if detected_skills else None

    existing_count = db.query(Resume).filter(
        Resume.user_id == current_user.id,
        Resume.filename == file.filename,
    ).count()

    content_hash = _compute_content_hash(text)

    new_resume = Resume(
        user_id=current_user.id,
        filename=file.filename,
        file_path=str(file_path),
        extracted_text=text,
        skills=skills_str,
        content_hash=content_hash,
        version=existing_count + 1,
        is_active=True,
        parsed_name=structured_fields.get("name"),
        parsed_email=structured_fields.get("email"),
        parsed_phone=structured_fields.get("phone"),
        parsed_location=structured_fields.get("location"),
        parsed_linkedin=structured_fields.get("linkedin"),
        parsed_github=structured_fields.get("github"),
        parsed_portfolio=structured_fields.get("portfolio"),
    )

    db.query(Resume).filter(
        Resume.user_id == current_user.id,
        Resume.is_active == True,
    ).update({"is_active": False})

    db.add(new_resume)
    db.commit()
    db.refresh(new_resume)

    version_record = ResumeVersion(
        resume_id=new_resume.id,
        user_id=current_user.id,
        version_number=1,
        filename=file.filename,
        file_path=str(file_path),
        extracted_text=text,
        skills=skills_str,
        content_hash=content_hash,
        change_reason="Initial upload",
    )
    db.add(version_record)
    db.commit()

    try:
        from app.ml.prediction.classifier_service import ClassifierService
        from app.ml.prediction.search_service import SearchService
        from app.ml.prediction.recommender_service import RecommenderService
        classification = ClassifierService.classify(text)
        SearchService.index_resume(new_resume.id, text)
        quality = RecommenderService.predict_quality(text)
    except Exception:
        classification = None
        quality = None

    duplicate_of = None
    existing = db.query(Resume).filter(
        Resume.user_id == current_user.id,
        Resume.content_hash == content_hash,
        Resume.id != new_resume.id,
    ).first()
    if existing:
        duplicate_of = existing.id

    return {
        "id": new_resume.id,
        "filename": new_resume.filename,
        "skills": detected_skills,
        "version": 1,
        "is_active": True,
        "parsed_fields": structured_fields,
        "classification": classification,
        "quality": quality,
        "duplicate_of": duplicate_of,
        "message": "Resume uploaded and skills extracted successfully.",
    }


@router.get("/", response_model=List[ResumeResponse])
def get_my_resumes(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    search: Optional[str] = Query(None),
    active_only: bool = Query(False),
):
    query = db.query(Resume).filter(Resume.user_id == current_user.id)

    if active_only:
        query = query.filter(Resume.is_active == True)

    if search:
        search_term = f"%{search}%"
        query = query.filter(
            (Resume.filename.ilike(search_term))
            | (Resume.skills.ilike(search_term))
            | (Resume.parsed_name.ilike(search_term))
        )

    total = query.count()
    resumes = (
        query.order_by(Resume.created_at.desc())
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )

    return resumes


@router.get("/active")
def get_active_resume(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.user_id == current_user.id,
        Resume.is_active == True,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="No active resume found")
    return ResumeResponse.model_validate(resume)


@router.get("/{resume_id}", response_model=ResumeResponse)
def get_resume(
    resume_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return resume


@router.get("/{resume_id}/versions")
def get_resume_versions(
    resume_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    versions = (
        db.query(ResumeVersion)
        .filter(
            ResumeVersion.resume_id == resume_id,
            ResumeVersion.user_id == current_user.id,
        )
        .order_by(ResumeVersion.version_number.desc())
        .all()
    )

    return [
        {
            "id": v.id,
            "version_number": v.version_number,
            "filename": v.filename,
            "content_hash": v.content_hash,
            "change_reason": v.change_reason,
            "created_at": v.created_at.isoformat() if v.created_at else None,
        }
        for v in versions
    ]


@router.post("/{resume_id}/versions")
def create_resume_version(
    resume_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    latest_version = (
        db.query(ResumeVersion)
        .filter(ResumeVersion.resume_id == resume_id)
        .order_by(ResumeVersion.version_number.desc())
        .first()
    )
    next_version = (latest_version.version_number + 1) if latest_version else 1

    version_record = ResumeVersion(
        resume_id=resume_id,
        user_id=current_user.id,
        version_number=next_version,
        filename=resume.filename,
        file_path=resume.file_path,
        extracted_text=resume.extracted_text,
        skills=resume.skills,
        content_hash=resume.content_hash,
        change_reason="Manual version snapshot",
    )
    db.add(version_record)
    db.commit()
    db.refresh(version_record)

    return {
        "id": version_record.id,
        "version_number": version_record.version_number,
        "message": f"Version {next_version} created.",
    }


class CompareRequest(BaseModel):
    resume_id_a: int
    resume_id_b: int


@router.post("/compare")
def compare_resumes(
    req: CompareRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume_a = db.query(Resume).filter(
        Resume.id == req.resume_id_a,
        Resume.user_id == current_user.id,
    ).first()
    resume_b = db.query(Resume).filter(
        Resume.id == req.resume_id_b,
        Resume.user_id == current_user.id,
    ).first()

    if not resume_a or not resume_b:
        raise HTTPException(status_code=404, detail="One or both resumes not found")

    skills_a = set(resume_a.skills.split(",")) if resume_a.skills else set()
    skills_b = set(resume_b.skills.split(",")) if resume_b.skills else set()
    skills_a = {s.strip() for s in skills_a if s.strip()}
    skills_b = {s.strip() for s in skills_b if s.strip()}

    common = skills_a & skills_b
    only_a = skills_a - skills_b
    only_b = skills_b - skills_a

    text_a = resume_a.extracted_text or ""
    text_b = resume_b.extracted_text or ""

    from app.ml.cache import ml_cache
    from app.ml.embeddings.embedder import Embedder
    import numpy as np

    try:
        emb_a = Embedder.encode(text_a[:5000])
        emb_b = Embedder.encode(text_b[:5000])
        similarity = float(Embedder.cosine_similarities(
            emb_a.reshape(1, -1),
            emb_b.reshape(1, -1)
        )[0][0])
    except Exception:
        similarity = 0.0

    word_count_a = len(text_a.split())
    word_count_b = len(text_b.split())

    return {
        "resume_a": {"id": resume_a.id, "filename": resume_a.filename, "word_count": word_count_a},
        "resume_b": {"id": resume_b.id, "filename": resume_b.filename, "word_count": word_count_b},
        "skills_common": sorted(common),
        "skills_only_a": sorted(only_a),
        "skills_only_b": sorted(only_b),
        "skill_overlap_pct": round(len(common) / max(len(skills_a | skills_b), 1) * 100, 1),
        "text_similarity": round(similarity, 4),
        "length_diff": word_count_a - word_count_b,
    }


@router.put("/{resume_id}/active")
def set_active_resume(
    resume_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    db.query(Resume).filter(
        Resume.user_id == current_user.id,
        Resume.is_active == True,
    ).update({"is_active": False})

    resume.is_active = True
    db.commit()

    return {"message": f"Resume '{resume.filename}' set as active.", "id": resume.id}


@router.delete("/{resume_id}")
def delete_resume(
    resume_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    if resume.file_path and os.path.exists(resume.file_path):
        os.remove(resume.file_path)

    try:
        from app.ml.prediction.search_service import SearchService
        SearchService.remove_resume(resume_id)
    except Exception:
        pass

    from app.models.career import (
        ResumeAnalysis,
        SkillGapAnalysis,
        OptimizedResume,
        CareerReadiness,
    )
    from app.models.ats_report import ATSReport
    from app.models.ml_analytics import (
        MLClassification,
        MLATSPrediction,
        MLSkillExtraction,
        MLJobRecommendation,
        MLQualityPrediction,
        MLAnalysisHistory,
    )

    db.query(ResumeAnalysis).filter(ResumeAnalysis.resume_id == resume_id).delete(
        synchronize_session=False
    )
    db.query(ATSReport).filter(ATSReport.resume_id == resume_id).delete(
        synchronize_session=False
    )
    db.query(MLClassification).filter(
        MLClassification.resume_id == resume_id
    ).delete(synchronize_session=False)
    db.query(MLATSPrediction).filter(
        MLATSPrediction.resume_id == resume_id
    ).delete(synchronize_session=False)
    db.query(MLSkillExtraction).filter(
        MLSkillExtraction.resume_id == resume_id
    ).delete(synchronize_session=False)
    db.query(MLJobRecommendation).filter(
        MLJobRecommendation.resume_id == resume_id
    ).delete(synchronize_session=False)
    db.query(MLQualityPrediction).filter(
        MLQualityPrediction.resume_id == resume_id
    ).delete(synchronize_session=False)
    db.query(MLAnalysisHistory).filter(
        MLAnalysisHistory.resume_id == resume_id
    ).delete(synchronize_session=False)
    db.query(ResumeVersion).filter(
        ResumeVersion.resume_id == resume_id
    ).delete(synchronize_session=False)

    analysis_ids = [
        a.id
        for a in db.query(ResumeAnalysis.id)
        .filter(ResumeAnalysis.resume_id == resume_id)
        .all()
    ]
    if analysis_ids:
        db.query(SkillGapAnalysis).filter(
            SkillGapAnalysis.resume_analysis_id.in_(analysis_ids)
        ).delete(synchronize_session=False)
        db.query(OptimizedResume).filter(
            OptimizedResume.resume_analysis_id.in_(analysis_ids)
        ).delete(synchronize_session=False)
        db.query(CareerReadiness).filter(
            CareerReadiness.resume_analysis_id.in_(analysis_ids)
        ).delete(synchronize_session=False)

    db.delete(resume)
    db.commit()

    return {"message": "Resume and all related records deleted."}


class BulkDeleteRequest(BaseModel):
    resume_ids: list[int]


@router.post("/bulk-delete")
def bulk_delete_resumes(
    req: BulkDeleteRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resumes = db.query(Resume).filter(
        Resume.id.in_(req.resume_ids),
        Resume.user_id == current_user.id,
    ).all()

    if not resumes:
        raise HTTPException(status_code=404, detail="No resumes found")

    deleted_count = 0
    for resume in resumes:
        rid = resume.id
        try:
            db.query(ResumeAnalysis).filter(ResumeAnalysis.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(ATSReport).filter(ATSReport.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(MLClassification).filter(MLClassification.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(MLATSPrediction).filter(MLATSPrediction.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(MLSkillExtraction).filter(MLSkillExtraction.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(MLJobRecommendation).filter(MLJobRecommendation.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(MLQualityPrediction).filter(MLQualityPrediction.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(MLAnalysisHistory).filter(MLAnalysisHistory.resume_id == rid).delete(
                synchronize_session=False
            )
            db.query(ResumeVersion).filter(ResumeVersion.resume_id == rid).delete(
                synchronize_session=False
            )
            analysis_ids = [
                a.id
                for a in db.query(ResumeAnalysis.id)
                .filter(ResumeAnalysis.resume_id == rid)
                .all()
            ]
            if analysis_ids:
                db.query(SkillGapAnalysis).filter(
                    SkillGapAnalysis.resume_analysis_id.in_(analysis_ids)
                ).delete(synchronize_session=False)
                db.query(OptimizedResume).filter(
                    OptimizedResume.resume_analysis_id.in_(analysis_ids)
                ).delete(synchronize_session=False)
                db.query(CareerReadiness).filter(
                    CareerReadiness.resume_analysis_id.in_(analysis_ids)
                ).delete(synchronize_session=False)
            try:
                from app.ml.prediction.search_service import SearchService
                SearchService.remove_resume(rid)
            except Exception:
                pass
            db.delete(resume)
            deleted_count += 1
        except Exception:
            continue

    db.commit()

    return {"message": f"{deleted_count} resume(s) deleted.", "deleted": deleted_count}
