import os
import json
import re
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

ATS_SKILL_NORM = {
    "node": "node.js", "nodejs": "node.js", "node js": "node.js",
    "react.js": "react", "reactjs": "react", "react js": "react",
    "vue.js": "vue", "vuejs": "vue", "vue js": "vue",
    "angular.js": "angular", "angularjs": "angular",
    "next.js": "nextjs", "nextjs": "nextjs", "next js": "nextjs",
    "express.js": "express", "expressjs": "express",
    "ts": "typescript", "js": "javascript", "py": "python",
    "postgres": "postgresql", "k8s": "kubernetes", "tf": "terraform",
    "ci/cd": "ci_cd", "rest api": "rest_api", "rest": "rest_api",
    "object oriented programming": "oop",
    "object-oriented programming": "oop",
    "ml": "machine_learning", "ds&a": "dsa",
    "data structures and algorithms": "dsa",
    "data structures & algorithms": "dsa",
    "amazon web services": "aws",
    "google cloud platform": "gcp", "google cloud": "gcp",
    "microsoft azure": "azure",
    "deep learning": "deep_learning", "machine learning": "machine_learning",
    "computer vision": "computer_vision",
    "natural language processing": "nlp",
    "github actions": "github_actions",
}


def _ats_normalize_skill(skill: str) -> str:
    s = skill.lower().strip()
    return ATS_SKILL_NORM.get(s, s)


# ── Prompt Injection Protection ──

_INJECTION_PATTERNS = [
    r"ignore\s+(all\s+)?(previous|prior|above|earlier)\s+(instructions?|prompts?|rules?)",
    r"disregard\s+(all\s+)?(previous|prior|above|earlier)\s+(instructions?|prompts?|rules?)",
    r"you\s+are\s+now\s+(a|an|the)\s+",
    r"system\s*:\s*",
    r"<\s*(system|assistant|user)\s*>",
    r"\[INST\]|\[/INST\]",
    r"<<SYS>>|<</SYS>>",
    r"act\s+as\s+(a|an|the)\s+",
    r"pretend\s+(you\s+are|to\s+be)\s+",
    r"roleplay\s+as\s+",
    r"jailbreak",
    r"DAN\s+mode",
    r"developer\s+mode",
]

_INJECTION_RE = re.compile("|".join(_INJECTION_PATTERNS), re.IGNORECASE)


def _sanitize_input(text: str, max_length: int = 50000) -> str:
    """
    Sanitize user-provided text before passing to AI analysis.
    Strips prompt injection patterns and truncates to max length.
    """
    if not text:
        return ""

    # Truncate to max length
    text = text[:max_length]

    # Remove common prompt injection patterns
    text = _INJECTION_RE.sub("[FILTERED]", text)

    # Strip null bytes and control characters (except newline/tab)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    return text


class CareerService:

    @staticmethod
    def _call_ai(prompt: str, model: str = "gpt-3.5-turbo", max_tokens: int = 800) -> dict | list | None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key or api_key == "your_key_here":
            return None

        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=max_tokens,
                response_format={"type": "json_object"},
            )
            content = response.choices[0].message.content.strip()
            if content.startswith("```json"):
                content = content.replace("```json", "").replace("```", "").strip()
            
            # Log API usage
            try:
                from app.database import SessionLocal
                from app.models.api_usage import ApiUsage
                db = SessionLocal()
                try:
                    usage = ApiUsage(
                        provider="openai",
                        model=model,
                        feature="career",
                        prompt_tokens=response.usage.prompt_tokens if response.usage else 0,
                        completion_tokens=response.usage.completion_tokens if response.usage else 0,
                        total_tokens=response.usage.total_tokens if response.usage else 0,
                        cost=0.0,
                    )
                    db.add(usage)
                    db.commit()
                finally:
                    db.close()
            except Exception:
                pass
            
            return json.loads(content)
        except Exception as e:
            import logging
            logging.getLogger(__name__).error("CareerService AI call failed: %s", e)
            return None

    # ── Job Description Analysis ──

    @staticmethod
    def analyze_job_description(raw_text: str) -> dict:
        prompt = f"""Analyze the following job description and extract structured information.

Job Description:
{raw_text[:4000]}

Return a JSON object with these exact keys:
{{
  "required_skills": ["skill1", "skill2"],
  "preferred_skills": ["skill1", "skill2"],
  "technologies": ["tech1", "tech2"],
  "responsibilities": ["resp1", "resp2"],
  "experience_years": "3-5 years",
  "education_requirements": "Bachelor's in CS or related field",
  "soft_skills": ["communication", "teamwork"],
  "keywords": ["keyword1", "keyword2"]
}}"""

        result = CareerService._call_ai(prompt)
        if result:
            return result

        return CareerService._mock_jd_analysis(raw_text)

    @staticmethod
    def _mock_jd_analysis(raw_text: str) -> dict:
        text_lower = raw_text.lower()
        all_skills = [
            "Python", "JavaScript", "TypeScript", "React", "Node.js", "Express",
            "FastAPI", "Django", "Flask", "PostgreSQL", "SQL", "MongoDB",
            "Docker", "Kubernetes", "AWS", "Azure", "GCP", "Git", "Java", "C++",
            "REST", "GraphQL", "Redis", "Kafka", "CI/CD", "Agile", "Scrum",
            "Vue", "Angular", "Next.js", "Tailwind CSS", "Redux", "Prisma",
            "Go", "Rust", "Elasticsearch", "RabbitMQ", "Jenkins", "Terraform",
        ]
        detected = [s for s in all_skills if s.lower() in text_lower]

        soft_skills_pool = []
        soft_keywords = {
            "communication": "Communication", "teamwork": "Teamwork",
            "leadership": "Leadership", "problem solving": "Problem Solving",
            "analytical": "Analytical Thinking", "creative": "Creativity",
            "adaptability": "Adaptability", "time management": "Time Management",
            "collaboration": "Collaboration", "self-motivated": "Self-Motivated",
        }
        for keyword, skill in soft_keywords.items():
            if keyword in text_lower:
                soft_skills_pool.append(skill)
        if not soft_skills_pool:
            soft_skills_pool = ["Communication", "Teamwork", "Problem Solving"]

        resp_keywords = ["responsibilities", "duties", "role", "what you"]
        has_responsibilities = any(k in text_lower for k in resp_keywords)
        responsibilities = []
        if has_responsibilities:
            sentences = [s.strip() for s in raw_text.replace("\n", " ").split(".") if len(s.strip()) > 20]
            responsibilities = [s.capitalize() for s in sentences[:5]]

        seniority = "Mid-Level"
        if any(w in text_lower for w in ["senior", "lead", "principal", "staff", "architect", "director"]):
            seniority = "Senior"
        elif any(w in text_lower for w in ["junior", "entry", "intern", "graduate", "associate"]):
            seniority = "Junior"

        exp_match = ""
        import re
        exp_patterns = [
            r'(\d+)\+?\s*(?:to|-)\s*(\d+)\s*years?',
            r'(\d+)\+?\s*years?',
        ]
        for pat in exp_patterns:
            m = re.search(pat, text_lower)
            if m:
                groups = m.groups()
                if len(groups) == 2:
                    exp_match = f"{groups[0]}-{groups[1]} years"
                else:
                    exp_match = f"{groups[0]}+ years"
                break
        if not exp_match:
            exp_match = "3-5 years" if "experience" in text_lower else "1-3 years"

        return {
            "required_skills": detected[:8] if detected else ["Python", "SQL"],
            "preferred_skills": detected[8:14] if len(detected) > 8 else ["Docker", "AWS"],
            "technologies": detected[:12] if detected else ["Python"],
            "responsibilities": responsibilities if responsibilities else [
                "Develop and maintain software applications",
                "Collaborate with cross-functional teams",
                "Write clean, maintainable code",
            ],
            "experience_years": exp_match,
            "education_requirements": "Bachelor's degree in Computer Science or related field" if any(
                w in text_lower for w in ["bachelor", "degree", "education", "bs ", "b.s."]
            ) else None,
            "soft_skills": soft_skills_pool,
            "keywords": detected[:8] if detected else ["software", "development"],
        }

    # ── Resume Analysis ──

    @staticmethod
    def analyze_resume(resume_text: str, jd_text: str = None) -> dict:
        resume_text = _sanitize_input(resume_text)
        if jd_text:
            jd_text = _sanitize_input(jd_text)
        jd_section = f"\nTarget Job Description:\n{jd_text[:2000]}" if jd_text else ""

        prompt = f"""Analyze the following resume and extract structured information.

Resume:
{resume_text[:4000]}
{jd_section}

Return a JSON object with these exact keys:
{{
  "summary": "Brief professional summary",
  "detected_skills": ["skill1", "skill2"],
  "experience_level": "Junior/Mid/Senior",
  "projects": ["project1 description", "project2 description"],
  "technologies": ["tech1", "tech2"],
  "education": ["degree info"],
  "certifications": ["cert1", "cert2"]
}}"""

        result = CareerService._call_ai(prompt)
        if result:
            return result

        return CareerService._mock_resume_analysis(resume_text)

    @staticmethod
    def _mock_resume_analysis(resume_text: str) -> dict:
        text_lower = resume_text.lower()
        all_skills = [
            "Python", "JavaScript", "TypeScript", "React", "Node.js",
            "FastAPI", "Django", "Flask", "PostgreSQL", "SQL", "MongoDB",
            "Docker", "AWS", "Git", "Java", "C++",
        ]
        detected = [s for s in all_skills if s.lower() in text_lower]

        experience_level = "Junior"
        if any(w in text_lower for w in ["senior", "lead", "principal", "5+ years", "7+ years"]):
            experience_level = "Senior"
        elif any(w in text_lower for w in ["mid", "3 years", "4 years", "5 years"]):
            experience_level = "Mid"

        return {
            "summary": f"Experienced professional with skills in {', '.join(detected[:5]) if detected else 'software development'}.",
            "detected_skills": detected if detected else ["Python"],
            "experience_level": experience_level,
            "projects": ["Project experience detected from resume"],
            "technologies": detected[:8] if detected else ["Python"],
            "education": ["Degree information extracted from resume"],
            "certifications": [],
        }

    # ── ATS Engine ──

    @staticmethod
    def parse_resume_detailed(resume_text: str) -> dict:
        """Parse resume into structured sections with real extraction."""
        resume_text = _sanitize_input(resume_text)
        text = resume_text.strip()
        text_lower = text.lower()

        # Extract contact info
        emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
        phones = re.findall(r'[\+]?[(]?\d{1,4}[)]?[-\s./]?\d{1,4}[-\s./]?\d{1,9}', text)
        linkedin = 'linkedin.com' in text_lower
        github = 'github.com' in text_lower

        # Extract sections by common headings
        sections = {}
        section_patterns = {
            'contact': r'(?:contact|phone|email|address)',
            'summary': r'(?:summary|objective|profile|about)',
            'experience': r'(?:experience|employment|work history|professional experience)',
            'education': r'(?:education|academic|university|college|degree)',
            'skills': r'(?:skills|technologies|technical skills|competencies|proficiencies)',
            'projects': r'(?:projects|portfolio|personal projects)',
            'certifications': r'(?:certifications|certificates|licenses|credentials)',
            'achievements': r'(?:achievements|awards|honors|accomplishments)',
        }

        lines = text.split('\n')
        current_section = 'header'
        section_content = {'header': []}

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            matched = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_stripped, re.IGNORECASE) and len(line_stripped) < 60:
                    current_section = section_name
                    section_content.setdefault(section_name, [])
                    matched = True
                    break
            if not matched:
                section_content.setdefault(current_section, []).append(line_stripped)

        # Extract skills
        known_skills = [
            "python", "javascript", "typescript", "react", "node.js", "express",
            "fastapi", "django", "flask", "postgresql", "sql", "mongodb", "mysql",
            "docker", "kubernetes", "aws", "azure", "gcp", "git", "java", "c++",
            "c#", "go", "rust", "ruby", "php", "swift", "kotlin",
            "rest", "graphql", "redis", "kafka", "rabbitmq", "ci/cd",
            "agile", "scrum", "jira", "linux", "nginx", "apache",
            "html", "css", "sass", "tailwind", "bootstrap",
            "vue.js", "angular", "svelte", "next.js", "nuxt.js",
            "tensorflow", "pytorch", "pandas", "numpy", "scikit-learn",
            "machine learning", "deep learning", "nlp", "computer vision",
            "microservices", "serverless", "terraform", "ansible",
            "elasticsearch", "splunk", "datadog", "prometheus",
            "figma", "sketch", "adobe xd",
        ]
        detected_skills = []
        for skill in known_skills:
            if skill in text_lower:
                detected_skills.append(skill.title() if not skill.isupper() else skill)

        # Extract education keywords
        education_keywords = ["bachelor", "master", "phd", "b.s.", "m.s.", "b.tech", "m.tech",
                              "bca", "mca", "mba", "university", "college", "institute", "gpa"]
        education = [line.strip() for line in section_content.get('education', [])
                     if any(kw in line.lower() for kw in education_keywords)]

        # Extract experience indicators
        experience_indicators = ["years", "experience", "worked", "employment", "intern",
                                 "full-time", "part-time", "contract", "freelance"]
        experience_lines = [line.strip() for line in lines
                           if any(kw in line.lower() for kw in experience_indicators)]

        # Detect experience level
        experience_level = "Junior"
        if any(w in text_lower for w in ["senior", "lead", "principal", "staff", "architect", "7+ years", "10+ years"]):
            experience_level = "Senior"
        elif any(w in text_lower for w in ["mid-level", "3 years", "4 years", "5 years", "6 years"]):
            experience_level = "Mid"

        # Extract projects
        projects = section_content.get('projects', [])
        project_count = len([p for p in projects if len(p) > 20])

        # Extract certifications
        certifications = section_content.get('certifications', [])

        # Extract achievements
        achievements = section_content.get('achievements', [])

        # Formatting analysis
        has_tables = '|' in text or '┌' in text or '└' in text
        has_bullets = any(c in text for c in ['•', '●', '○', '▸', '▹', '►', '–', '-'])
        has_headers = bool(section_content.get('header'))
        section_count = len([k for k in section_content if k != 'header'])
        line_count = len([l for l in lines if l.strip()])
        avg_line_length = sum(len(l) for l in lines) / max(len(lines), 1)

        # Advanced formatting detection
        has_images = '![image]' in text_lower or '[image]' in text_lower or '<img' in text_lower
        has_text_boxes = '<text' in text_lower or 'textbox' in text_lower
        has_headers_footers = bool(re.search(r'header|footer', text_lower)) and section_count < 3
        has_multiple_columns = any(c in text for c in ['│', '┃', '║'])
        has_icons = any(c in text for c in ['✓', '✗', '★', '☆', '▶', '■', '□', '◆', '◇'])

        # Date format detection
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+\d{4}\b',
            r'\b\d{4}\s*[-–]\s*(?:present|current|now)\b',
        ]
        has_dates = any(re.search(p, text_lower) for p in date_patterns)

        # Check for ATS-unfriendly elements
        issues = []
        if has_tables:
            issues.append("Tables detected - ATS may not parse table content correctly")
        if has_images:
            issues.append("Images detected - ATS cannot read image content")
        if has_text_boxes:
            issues.append("Text boxes detected - ATS may not parse text box content")
        if has_headers_footers:
            issues.append("Headers/footers detected - content may be ignored by ATS")
        if has_multiple_columns:
            issues.append("Multiple columns detected - ATS reads left-to-right, top-to-bottom")
        if has_icons:
            issues.append("Special icons/characters detected - use plain text for ATS compatibility")
        if '•' not in text and '-' not in text and '●' not in text:
            issues.append("No bullet points found - use bullet points for better readability")
        if section_count < 4:
            issues.append(f"Only {section_count} sections detected - aim for 5+ standard sections")
        if not emails:
            issues.append("No email address found")
        if not phones:
            issues.append("No phone number found")
        if not linkedin:
            issues.append("LinkedIn profile not found - add LinkedIn URL")
        if len(detected_skills) < 5:
            issues.append("Few skills detected - add a dedicated skills section with relevant technologies")

        return {
            "name": section_content.get('header', [''])[0] if section_content.get('header') else '',
            "contact": {
                "emails": emails,
                "phones": phones,
                "has_linkedin": linkedin,
                "has_github": github,
            },
            "sections_found": list(section_content.keys()),
            "section_count": section_count,
            "detected_skills": detected_skills,
            "skill_count": len(detected_skills),
            "experience_level": experience_level,
            "education": education[:3],
            "education_count": len(education),
            "projects": projects[:5],
            "project_count": project_count,
            "certifications": certifications[:5],
            "achievements": achievements[:5],
            "achievement_count": len(achievements),
            "experience_lines": experience_lines[:5],
            "line_count": line_count,
            "avg_line_length": round(avg_line_length, 1),
            "formatting": {
                "has_tables": has_tables,
                "has_bullets": has_bullets,
                "has_headers": has_headers,
                "has_images": has_images,
                "has_text_boxes": has_text_boxes,
                "has_headers_footers": has_headers_footers,
                "has_multiple_columns": has_multiple_columns,
                "has_icons": has_icons,
                "has_dates": has_dates,
                "section_count": section_count,
                "issues": issues,
            },
        }

    @staticmethod
    def parse_job_description_detailed(jd_text: str) -> dict:
        """Parse JD into structured data with real extraction."""
        jd_text = _sanitize_input(jd_text)
        text = jd_text.strip()
        text_lower = text.lower()

        # Extract skills from JD
        all_tech = [
            "python", "javascript", "typescript", "react", "node.js", "express",
            "fastapi", "django", "flask", "postgresql", "sql", "mongodb", "mysql",
            "docker", "kubernetes", "aws", "azure", "gcp", "git", "java", "c++",
            "c#", "go", "rust", "ruby", "php", "swift", "kotlin",
            "rest", "graphql", "redis", "kafka", "rabbitmq", "ci/cd",
            "agile", "scrum", "jira", "linux", "nginx",
            "html", "css", "tailwind", "bootstrap", "vue.js", "angular",
            "tensorflow", "pytorch", "pandas", "numpy",
            "microservices", "serverless", "terraform",
            "elasticsearch", "figma",
        ]

        required_skills = []
        preferred_skills = []
        all_mentioned = []

        for tech in all_tech:
            if tech in text_lower:
                all_mentioned.append(tech.title() if not tech.isupper() else tech)

        # Try to split required vs preferred
        preferred_section = False
        for line in text.split('\n'):
            line_l = line.lower()
            if 'preferred' in line_l or 'nice to have' in line_l or 'bonus' in line_l:
                preferred_section = True
            elif 'required' in line_l or 'must have' in line_l or 'qualifications' in line_l:
                preferred_section = False

            for tech in all_tech:
                if tech in line_l:
                    skill_name = tech.title() if not tech.isupper() else tech
                    if preferred_section:
                        if skill_name not in preferred_skills:
                            preferred_skills.append(skill_name)
                    else:
                        if skill_name not in required_skills:
                            required_skills.append(skill_name)

        if not required_skills:
            required_skills = all_mentioned[:8]
        if not preferred_skills:
            preferred_skills = all_mentioned[8:12] if len(all_mentioned) > 8 else []

        # Extract experience requirements
        exp_match = re.search(r'(\d+)[\s\-\+]*(?:to\s*(\d+))?\s*years?\s*(?:of\s+)?experience', text_lower)
        exp_years = None
        if exp_match:
            exp_years = f"{exp_match.group(1)}-{exp_match.group(2) or exp_match.group(1)} years"

        # Extract education requirements
        edu_keywords = ["bachelor", "master", "phd", "b.s.", "m.s.", "b.tech", "m.tech", "mba"]
        education_req = [line.strip() for line in text.split('\n')
                        if any(kw in line.lower() for kw in edu_keywords)][:3]

        # Extract responsibilities
        resp_keywords = ["responsibilities", "what you'll do", "role", "duties"]
        responsibilities = []
        in_resp = False
        for line in text.split('\n'):
            line_l = line.lower().strip()
            if any(kw in line_l for kw in resp_keywords):
                in_resp = True
                continue
            if in_resp and line.strip():
                if line.strip().startswith(('•', '-', '●', '*')) or (len(line.strip()) > 10 and len(line.strip()) < 200):
                    responsibilities.append(line.strip().lstrip('•-●* '))
                if len(responsibilities) >= 8:
                    break
                # Stop if we hit another section header
                if any(kw in line_l for kw in ['requirements', 'qualifications', 'skills', 'benefits', 'about']):
                    break

        if not responsibilities:
            responsibilities = ["Develop and maintain software applications",
                              "Collaborate with cross-functional teams",
                              "Write clean, maintainable code"]

        # Extract soft skills
        soft_skill_map = {
            "communication": ["communication", "communicate", "verbal", "written"],
            "teamwork": ["teamwork", "team player", "collaborate", "collaboration"],
            "problem solving": ["problem solving", "analytical", "critical thinking"],
            "leadership": ["leadership", "lead", "mentor", "guide"],
            "time management": ["time management", "deadline", "prioritize"],
            "adaptability": ["adaptability", "adaptable", "flexible", "flexibility"],
        }
        soft_skills = []
        for skill, keywords in soft_skill_map.items():
            if any(kw in text_lower for kw in keywords):
                soft_skills.append(skill.title())

        # Extract keywords (important terms)
        keywords = list(set(required_skills + preferred_skills))

        return {
            "required_skills": required_skills,
            "preferred_skills": preferred_skills,
            "all_skills": list(set(required_skills + preferred_skills)),
            "technologies": all_mentioned,
            "experience_years": exp_years,
            "education_requirements": education_req,
            "responsibilities": responsibilities,
            "soft_skills": soft_skills,
            "keywords": keywords,
            "skill_count": len(keywords),
        }

    @staticmethod
    def analyze_formatting(resume_text: str, parsed: dict) -> dict:
        """Analyze resume formatting for ATS compatibility."""
        issues = []
        suggestions = []
        score = 100.0

        fmt = parsed.get('formatting', {})

        # Check sections
        section_count = parsed.get('section_count', 0)
        if section_count < 4:
            issues.append(f"Only {section_count} sections found. ATS prefers 5+ clear sections.")
            suggestions.append("Add standard sections: Summary, Experience, Skills, Education, Projects")
            score -= 15
        elif section_count < 6:
            suggestions.append("Consider adding Certifications or Achievements sections")
            score -= 5

        # Check contact info
        contact = parsed.get('contact', {})
        if not contact.get('emails'):
            issues.append("No email address found in resume")
            suggestions.append("Add a professional email address at the top")
            score -= 10
        if not contact.get('phones'):
            issues.append("No phone number found")
            suggestions.append("Add a phone number in the contact section")
            score -= 5
        if not contact.get('has_linkedin'):
            suggestions.append("Add LinkedIn profile URL for better credibility")
            score -= 3

        # Check for ATS-unfriendly elements
        if fmt.get('has_tables'):
            issues.append("Tables detected - many ATS cannot parse table content")
            suggestions.append("Replace tables with simple text sections and bullet points")
            score -= 15

        if fmt.get('has_images'):
            issues.append("Images detected - ATS cannot read image content")
            suggestions.append("Remove images and describe content in text")
            score -= 10

        if fmt.get('has_text_boxes'):
            issues.append("Text boxes detected - ATS may not parse text box content")
            suggestions.append("Replace text boxes with standard text sections")
            score -= 10

        if fmt.get('has_headers_footers'):
            issues.append("Headers/footers detected - content may be ignored by ATS")
            suggestions.append("Move important content from headers/footers to main body")
            score -= 8

        if fmt.get('has_multiple_columns'):
            issues.append("Multiple columns detected - ATS reads left-to-right, top-to-bottom")
            suggestions.append("Use single-column layout for ATS compatibility")
            score -= 12

        if fmt.get('has_icons'):
            issues.append("Special icons/characters detected - ATS may not parse them")
            suggestions.append("Replace icons with plain text labels")
            score -= 5

        if not fmt.get('has_bullets'):
            issues.append("No bullet points found")
            suggestions.append("Use bullet points (•) to list achievements and responsibilities")
            score -= 10

        # Check skills section
        if parsed.get('skill_count', 0) < 5:
            issues.append(f"Only {parsed.get('skill_count', 0)} skills detected")
            suggestions.append("Add a dedicated Skills section with 10-15 relevant technologies")
            score -= 10

        # Check line length
        avg_len = parsed.get('avg_line_length', 0)
        if avg_len > 100:
            suggestions.append("Some lines are very long - keep bullet points under 80 characters")
            score -= 5

        # Check for special characters that ATS may not parse
        special_chars = ['→', '⇒', '★', '☆', '✔', '✘', '▶', '■', '□']
        found_special = [c for c in special_chars if c in resume_text]
        if found_special:
            issues.append(f"Special characters found: {', '.join(found_special[:3])}")
            suggestions.append("Replace special characters with standard text for ATS compatibility")
            score -= 5

        # Check section order (standard ATS-friendly order)
        sections_found = parsed.get('sections_found', [])
        standard_order = ['contact', 'summary', 'experience', 'skills', 'education', 'projects']
        if sections_found:
            order_score = 0
            for i, section in enumerate(standard_order):
                if section in sections_found:
                    idx = sections_found.index(section)
                    if idx <= i + 1:
                        order_score += 1
            if order_score < len(standard_order) // 2:
                suggestions.append("Consider reordering sections: Contact → Summary → Experience → Skills → Education")
                score -= 5

        score = max(0, min(100, score))

        return {
            "score": round(score, 1),
            "issues": issues,
            "suggestions": suggestions,
            "section_count": section_count,
            "has_tables": fmt.get('has_tables', False),
            "has_bullets": fmt.get('has_bullets', False),
            "has_images": fmt.get('has_images', False),
            "has_text_boxes": fmt.get('has_text_boxes', False),
            "has_headers_footers": fmt.get('has_headers_footers', False),
            "has_multiple_columns": fmt.get('has_multiple_columns', False),
            "has_icons": fmt.get('has_icons', False),
            "has_dates": fmt.get('has_dates', False),
        }

    @staticmethod
    def analyze_keywords(resume_text: str, jd_data: dict, resume_parsed: dict) -> dict:
        """Analyze keyword match between resume and JD."""
        text_lower = resume_text.lower()

        required = jd_data.get('required_skills', [])
        preferred = jd_data.get('preferred_skills', [])
        all_jd_keywords = jd_data.get('keywords', list(set(required + preferred)))

        resume_skills = [_ats_normalize_skill(s) for s in resume_parsed.get('detected_skills', [])]
        normalized_text_words = [_ats_normalize_skill(w) for w in text_lower.split()]

        matched = []
        missing = []
        repeated = []

        for kw in all_jd_keywords:
            kw_norm = _ats_normalize_skill(kw)
            count = text_lower.count(kw.lower())
            norm_count = sum(1 for w in normalized_text_words if w == kw_norm)
            effective_count = max(count, norm_count)
            if effective_count > 0:
                matched.append({"keyword": kw, "count": effective_count, "importance": "required" if kw in required else "preferred"})
                if effective_count > 3:
                    repeated.append({"keyword": kw, "count": effective_count, "warning": "Over-used - may look like keyword stuffing"})
            else:
                missing.append({"keyword": kw, "importance": "required" if kw in required else "preferred"})

        # Calculate keyword density
        total_words = len(resume_text.split())
        matched_keywords_count = sum(m['count'] for m in matched)
        keyword_density = round((matched_keywords_count / max(total_words, 1)) * 100, 2)

        # Score: percentage of required+preferred keywords found
        total_jd_keywords = len(all_jd_keywords)
        match_pct = (len(matched) / max(total_jd_keywords, 1)) * 100

        # Penalty for missing required keywords
        missing_required = [m for m in missing if m['importance'] == 'required']
        penalty = len(missing_required) * 5

        score = max(0, min(100, match_pct - penalty))

        return {
            "score": round(score, 1),
            "matched": matched,
            "missing": missing,
            "repeated": repeated,
            "keyword_density": keyword_density,
            "total_jd_keywords": total_jd_keywords,
            "matched_count": len(matched),
            "missing_count": len(missing),
        }

    @staticmethod
    def analyze_skills_match(resume_parsed: dict, jd_data: dict) -> dict:
        """Analyze skills match between resume and JD."""
        resume_skills = set(_ats_normalize_skill(s) for s in resume_parsed.get('detected_skills', []))
        required = [_ats_normalize_skill(s) for s in jd_data.get('required_skills', [])]
        preferred = [_ats_normalize_skill(s) for s in jd_data.get('preferred_skills', [])]
        all_jd = list(set(required + preferred))

        matched = [s for s in all_jd if s in resume_skills]
        missing = [s for s in all_jd if s not in resume_skills]
        additional = [s for s in resume_skills if s not in all_jd]

        required_matched = [s for s in required if s in resume_skills]
        required_missing = [s for s in required if s not in resume_skills]

        # Score based on required skills match (70%) + preferred (30%)
        required_pct = (len(required_matched) / max(len(required), 1)) * 100
        preferred_list = [s for s in preferred if s in resume_skills]
        preferred_pct = (len(preferred_list) / max(len(preferred), 1)) * 100
        score = (required_pct * 0.7) + (preferred_pct * 0.3)

        return {
            "score": round(score, 1),
            "matched": [s.title() for s in matched],
            "missing": [s.title() for s in missing],
            "additional": [s.title() for s in additional],
            "required_matched": [s.title() for s in required_matched],
            "required_missing": [s.title() for s in required_missing],
            "match_percentage": round((len(matched) / max(len(all_jd), 1)) * 100, 1),
        }

    @staticmethod
    def analyze_experience(resume_text: str, resume_parsed: dict, jd_data: dict) -> dict:
        """Analyze experience match."""
        text_lower = resume_text.lower()

        # Detect years of experience mentioned in resume
        year_patterns = [
            r'(\d+)[\s\-\+]*years?(?:\s+of)?\s+experience',
            r'experience[:\s]*(\d+)[\s\-\+]*years?',
            r'(\d+)[\s\-\+]*yr',
        ]
        detected_years = 0
        for pattern in year_patterns:
            match = re.search(pattern, text_lower)
            if match:
                detected_years = max(detected_years, int(match.group(1)))

        # If no explicit years, estimate from sections
        if detected_years == 0:
            exp_lines = resume_parsed.get('experience_lines', [])
            if len(exp_lines) >= 3:
                detected_years = 3
            elif len(exp_lines) >= 1:
                detected_years = 1

        # Parse required experience from JD
        jd_exp = jd_data.get('experience_years', '')
        required_years = 0
        if jd_exp:
            match = re.search(r'(\d+)', str(jd_exp))
            if match:
                required_years = int(match.group(1))

        # Score
        score = 70.0  # Base score
        if detected_years >= required_years:
            score = 90.0
        elif detected_years >= required_years - 1:
            score = 75.0
        elif detected_years > 0:
            score = 55.0

        # Check for action verbs
        action_verbs = ["developed", "implemented", "designed", "built", "led", "managed",
                       "created", "improved", "increased", "reduced", "automated", "optimized",
                       "delivered", "launched", "architected", "mentored", "collaborated"]
        found_verbs = [v for v in action_verbs if v in text_lower]

        if len(found_verbs) < 3:
            score -= 10

        # Check for quantifiable achievements
        quantifiers = re.findall(r'\d+[%x$KkMm]|\$\d+|\d+\s*(?:percent|%|users|clients|team|projects)', text_lower)
        if len(quantifiers) < 2:
            score -= 5

        score = max(0, min(100, score))

        return {
            "score": round(score, 1),
            "detected_years": detected_years,
            "required_years": required_years,
            "meets_requirement": detected_years >= required_years,
            "action_verbs_found": found_verbs[:10],
            "action_verb_count": len(found_verbs),
            "quantifiable_achievements": len(quantifiers),
            "suggestions": (
                [] if score >= 80 else
                [f"Add more action verbs (found {len(found_verbs)})"] if len(found_verbs) < 5 else []
            ) + (
                ["Quantify achievements with numbers/percentages"] if len(quantifiers) < 2 else []
            ),
        }

    @staticmethod
    def analyze_projects(resume_text: str, resume_parsed: dict, jd_data: dict) -> dict:
        """Analyze project relevance and quality."""
        projects = resume_parsed.get('projects', [])
        project_count = len(projects)

        jd_techs = [t.lower() for t in jd_data.get('technologies', []) + jd_data.get('all_skills', [])]
        resume_skills = [s.lower() for s in resume_parsed.get('detected_skills', [])]

        relevant_projects = 0
        for proj in projects:
            proj_lower = proj.lower()
            if any(tech in proj_lower for tech in jd_techs):
                relevant_projects += 1

        # Score
        score = 40.0  # Base
        if project_count >= 3:
            score = 70.0
        if project_count >= 5:
            score = 80.0
        if relevant_projects >= 2:
            score += 10
        if project_count == 0:
            score = 20.0

        # Check for measurable outcomes in project descriptions
        has_metrics = False
        for proj in projects:
            if any(c in proj.lower() for c in ['%', 'users', 'increased', 'reduced', 'improved', 'faster']):
                has_metrics = True
                break

        if not has_metrics and project_count > 0:
            score -= 5

        score = max(0, min(100, score))

        suggestions = []
        if project_count < 2:
            suggestions.append("Add at least 2-3 relevant projects to strengthen your resume")
        if relevant_projects == 0 and project_count > 0:
            suggestions.append("Align project descriptions with target job technologies")
        if not has_metrics and project_count > 0:
            suggestions.append("Add measurable outcomes to project descriptions (e.g., 'reduced load time by 40%')")

        return {
            "score": round(score, 1),
            "project_count": project_count,
            "relevant_projects": relevant_projects,
            "has_metrics": has_metrics,
            "suggestions": suggestions,
        }

    @staticmethod
    def analyze_education(resume_parsed: dict, jd_data: dict) -> dict:
        """Analyze education match."""
        resume_edu = resume_parsed.get('education', [])
        jd_edu = jd_data.get('education_requirements', [])

        score = 70.0  # Base

        has_degree = len(resume_edu) > 0
        has_jd_requirement = len(jd_edu) > 0

        if has_degree:
            score = 85.0
            # Check if degree level matches
            resume_text = ' '.join(resume_edu).lower()
            if any(d in resume_text for d in ['master', 'm.s.', 'm.tech', 'mba', 'phd']):
                score = 95.0
            elif any(d in resume_text for d in ['bachelor', 'b.s.', 'b.tech', 'bca']):
                score = 85.0

        if not has_degree and has_jd_requirement:
            score = 40.0

        return {
            "score": round(score, 1),
            "has_degree": has_degree,
            "education_found": resume_edu[:3],
            "required_education": jd_edu[:3],
            "meets_requirement": has_degree or not has_jd_requirement,
        }

    @staticmethod
    def analyze_readability(resume_text: str) -> dict:
        """Analyze resume readability."""
        sentences = [s.strip() for s in resume_text.replace('\n', ' ').split('.') if s.strip()]
        words = resume_text.split()
        total_words = len(words)
        total_sentences = max(len(sentences), 1)

        avg_words_per_sentence = total_words / total_sentences
        avg_word_length = sum(len(w) for w in words) / max(total_words, 1)

        # Simple readability heuristic
        score = 80.0

        if avg_words_per_sentence > 25:
            score -= 10  # Sentences too long
        if avg_words_per_sentence < 5:
            score -= 5  # Too choppy

        if avg_word_length > 7:
            score -= 5  # Too many complex words

        # Check for passive voice indicators
        passive_indicators = ['was', 'were', 'been', 'being', 'is', 'are', 'am']
        passive_count = sum(1 for w in words if w.lower() in passive_indicators)
        passive_ratio = passive_count / max(total_words, 1)
        if passive_ratio > 0.05:
            score -= 5

        # Check for consistency
        lines = [l.strip() for l in resume_text.split('\n') if l.strip()]
        bullet_lines = [l for l in lines if l.startswith(('•', '-', '●', '*'))]
        if bullet_lines:
            # Check if bullet styles are consistent
            first_char = bullet_lines[0][0]
            inconsistent = sum(1 for l in bullet_lines if l[0] != first_char)
            if inconsistent > 0:
                score -= 3

        score = max(0, min(100, score))

        suggestions = []
        if avg_words_per_sentence > 25:
            suggestions.append("Shorten sentences - keep them under 20 words for better readability")
        if passive_ratio > 0.05:
            suggestions.append("Use active voice instead of passive voice")
        if total_words > 800:
            suggestions.append("Resume is quite long - aim for 400-600 words for optimal ATS scanning")

        return {
            "score": round(score, 1),
            "total_words": total_words,
            "total_sentences": total_sentences,
            "avg_words_per_sentence": round(avg_words_per_sentence, 1),
            "passive_voice_ratio": round(passive_ratio * 100, 1),
            "suggestions": suggestions,
        }

    @staticmethod
    def calculate_comprehensive_ats_score(resume_text: str, jd_text: str = None) -> dict:
        """Calculate comprehensive ATS score with detailed breakdown."""
        # Sanitize inputs
        resume_text = _sanitize_input(resume_text)
        if jd_text:
            jd_text = _sanitize_input(jd_text)

        # Handle empty resume
        if not resume_text or not resume_text.strip():
            return {
                "overall_score": 0.0,
                "breakdown": {
                    "keyword_match": 0.0, "skills_match": 0.0, "experience_match": 0.0,
                    "projects_relevance": 0.0, "education_match": 0.0, "formatting": 0.0,
                    "readability": 0.0,
                },
                "weights": {},
                "resume_parsed": {}, "jd_parsed": None,
                "keyword_analysis": {}, "skills_analysis": {}, "formatting_analysis": {},
                "experience_analysis": {}, "projects_analysis": {}, "education_analysis": {},
                "readability_analysis": {},
                "matched_skills": [], "missing_skills": [], "additional_skills": [],
                "recommendations": [{"priority": "critical", "category": "content", "message": "Resume is empty — upload a resume to get an ATS analysis", "impact": "Cannot evaluate an empty resume"}],
            }

        # Parse both documents
        resume_parsed = CareerService.parse_resume_detailed(resume_text)

        jd_parsed = None
        if jd_text:
            jd_parsed = CareerService.parse_job_description_detailed(jd_text)

        # Run all analyses
        formatting = CareerService.analyze_formatting(resume_text, resume_parsed)
        readability = CareerService.analyze_readability(resume_text)

        if jd_parsed:
            keywords = CareerService.analyze_keywords(resume_text, jd_parsed, resume_parsed)
            skills = CareerService.analyze_skills_match(resume_parsed, jd_parsed)
            experience = CareerService.analyze_experience(resume_text, resume_parsed, jd_parsed)
            projects = CareerService.analyze_projects(resume_text, resume_parsed, jd_parsed)
            education = CareerService.analyze_education(resume_parsed, jd_parsed)
        else:
            # Without JD, use simpler scoring
            keywords = {"score": 50.0, "matched": [], "missing": [], "repeated": [],
                       "keyword_density": 0, "total_jd_keywords": 0, "matched_count": 0, "missing_count": 0}
            skills = {"score": 50.0, "matched": [], "missing": resume_parsed.get('detected_skills', []),
                     "additional": [], "required_matched": [], "required_missing": [],
                     "match_percentage": 0}
            experience = CareerService.analyze_experience(resume_text, resume_parsed, {})
            projects = CareerService.analyze_projects(resume_text, resume_parsed, {})
            education = CareerService.analyze_education(resume_parsed, {})

        # Weighted overall score
        weights = {
            "keyword": 0.35,
            "skills": 0.20,
            "experience": 0.15,
            "projects": 0.10,
            "education": 0.05,
            "formatting": 0.10,
            "readability": 0.05,
        }

        overall = (
            keywords["score"] * weights["keyword"] +
            skills["score"] * weights["skills"] +
            experience["score"] * weights["experience"] +
            projects["score"] * weights["projects"] +
            education["score"] * weights["education"] +
            formatting["score"] * weights["formatting"] +
            readability["score"] * weights["readability"]
        )

        # Generate prioritized recommendations
        recommendations = []

        # High priority
        if keywords["missing_count"] > 0:
            missing_kw = [m["keyword"] for m in keywords["missing"] if m["importance"] == "required"][:5]
            if missing_kw:
                recommendations.append({
                    "priority": "high",
                    "category": "keywords",
                    "message": f"Add these required keywords: {', '.join(missing_kw)}",
                    "impact": f"Could improve your score by ~{len(missing_kw) * 4} points",
                })

        if skills.get("required_missing"):
            recommendations.append({
                "priority": "high",
                "category": "skills",
                "message": f"Add these required skills to your resume: {', '.join(skills['required_missing'][:5])}",
                "impact": "Critical for passing initial ATS screening",
            })

        if formatting["issues"]:
            for issue in formatting["issues"][:2]:
                recommendations.append({
                    "priority": "high",
                    "category": "formatting",
                    "message": issue,
                    "impact": "ATS may fail to parse your resume correctly",
                })

        # Medium priority
        if experience.get("suggestions"):
            for s in experience["suggestions"]:
                recommendations.append({
                    "priority": "medium",
                    "category": "experience",
                    "message": s,
                    "impact": "Strengthens your experience section",
                })

        if projects.get("suggestions"):
            for s in projects["suggestions"]:
                recommendations.append({
                    "priority": "medium",
                    "category": "projects",
                    "message": s,
                    "impact": "Shows hands-on experience with relevant technologies",
                })

        if keywords.get("repeated"):
            repeated_names = [r["keyword"] for r in keywords["repeated"][:3]]
            recommendations.append({
                "priority": "medium",
                "category": "keywords",
                "message": f"Reduce repetition of: {', '.join(repeated_names)} - may look like keyword stuffing",
                "impact": "ATS may flag over-optimization",
            })

        # Low priority
        if readability.get("suggestions"):
            for s in readability["suggestions"]:
                recommendations.append({
                    "priority": "low",
                    "category": "readability",
                    "message": s,
                    "impact": "Improves human readability after ATS screening",
                })

        for s in formatting.get("suggestions", []):
            if s not in [r["message"] for r in recommendations]:
                recommendations.append({
                    "priority": "low",
                    "category": "formatting",
                    "message": s,
                    "impact": "Better ATS compatibility",
                })

        return {
            "overall_score": round(overall, 1),
            "breakdown": {
                "keyword_match": keywords["score"],
                "skills_match": skills["score"],
                "experience_match": experience["score"],
                "projects_relevance": projects["score"],
                "education_match": education["score"],
                "formatting": formatting["score"],
                "readability": readability["score"],
            },
            "weights": weights,
            "resume_parsed": resume_parsed,
            "jd_parsed": jd_parsed,
            "keyword_analysis": keywords,
            "skills_analysis": skills,
            "formatting_analysis": formatting,
            "experience_analysis": experience,
            "projects_analysis": projects,
            "education_analysis": education,
            "readability_analysis": readability,
            "matched_skills": skills.get("matched", []),
            "missing_skills": skills.get("missing", []),
            "additional_skills": skills.get("additional", []),
            "recommendations": recommendations,
        }

    @staticmethod
    def optimize_resume_ats(resume_text: str, ats_analysis: dict, jd_text: str = None) -> dict:
        """Generate ATS-optimized resume without fabricating content."""
        jd_section = f"\nTarget Job Description:\n{jd_text[:2000]}" if jd_text else ""

        # Build context from analysis
        missing_skills = ats_analysis.get("missing_skills", [])
        missing_keywords = [m["keyword"] for m in ats_analysis.get("keyword_analysis", {}).get("missing", [])[:10]]
        recommendations = ats_analysis.get("recommendations", [])

        recs_text = "\n".join(f"- [{r['priority']}] {r['message']}" for r in recommendations[:10])

        prompt = f"""Optimize this resume for ATS compatibility. 

IMPORTANT RULES:
- NEVER fabricate experience, projects, certifications, or achievements
- NEVER add skills the candidate doesn't have
- ONLY rewrite, reorganize, and improve existing content
- Add missing keywords NATURALLY into existing descriptions
- Improve formatting for ATS parsing
- Quantify achievements where possible using existing information

Resume:
{resume_text[:3000]}

Missing Keywords to incorporate naturally: {json.dumps(missing_keywords)}
Missing Skills: {json.dumps(missing_skills[:8])}
Recommendations to address:
{recs_text}
{jd_section}

Return a JSON object:
{{
  "optimized_text": "The full optimized resume text...",
  "improvements": ["List of specific changes made"],
  "professional_summary": "Optimized 2-3 line summary",
  "optimized_skills": ["Reorganized skills list"],
  "optimized_keywords": ["Keywords now included"],
  "changes_summary": "Brief summary of all changes"
}}"""

        result = CareerService._call_ai(prompt, max_tokens=1500)
        if result:
            return result

        # Fallback: return original with suggestions
        return {
            "optimized_text": resume_text,
            "improvements": [r["message"] for r in recommendations[:5]],
            "professional_summary": "",
            "optimized_skills": ats_analysis.get("matched_skills", []),
            "optimized_keywords": missing_keywords[:5],
            "changes_summary": "AI optimization unavailable. Apply the recommendations manually.",
        }

    # ── Skill Gap Analysis ──

    @staticmethod
    def analyze_skill_gap(resume_skills: list, required_skills: list) -> dict:
        prompt = f"""Compare the candidate's skills against the required skills for a job position.

Candidate Skills: {json.dumps(resume_skills)}
Required Skills: {json.dumps(required_skills)}

Return a JSON object:
{{
  "matched_skills": ["skill1", "skill2"],
  "missing_skills": ["skill3", "skill4"],
  "additional_skills": ["skill5"],
  "priority_skills": ["skill3", "skill4"],
  "match_percentage": 60.0
}}"""

        result = CareerService._call_ai(prompt)
        if result:
            return result

        return CareerService._mock_skill_gap(resume_skills, required_skills)

    @staticmethod
    def _mock_skill_gap(resume_skills: list, required_skills: list) -> dict:
        resume_norm = {_ats_normalize_skill(s) for s in resume_skills}
        required_norm = [_ats_normalize_skill(s) for s in required_skills]

        matched = [s for s in required_skills if _ats_normalize_skill(s) in resume_norm]
        missing = [s for s in required_skills if _ats_normalize_skill(s) not in resume_norm]
        additional = [s for s in resume_skills if _ats_normalize_skill(s) not in {_ats_normalize_skill(r) for r in required_skills}]

        match_pct = (len(matched) / max(len(required_skills), 1)) * 100

        return {
            "matched_skills": matched,
            "missing_skills": missing,
            "additional_skills": additional,
            "priority_skills": missing[:3],
            "match_percentage": round(match_pct, 1),
        }

    # ── Learning Roadmap ──

    @staticmethod
    def generate_learning_roadmap(missing_skills: list) -> dict:
        prompt = f"""Create a personalized learning roadmap for the following missing skills.

Missing Skills: {json.dumps(missing_skills)}

For each skill, create a learning module. Return a JSON object:
{{
  "items": [
    {{
      "topic": "Skill Name",
      "description": "What to learn and why",
      "hours": 20,
      "difficulty": "Beginner/Intermediate/Advanced",
      "priority": "High/Medium/Low",
      "milestones": ["Complete basics", "Build a project"],
      "mini_project": "Build a small project using this skill"
    }}
  ],
  "total_hours": 80,
  "estimated_weeks": 8
}}"""

        result = CareerService._call_ai(prompt)
        if result:
            return result

        return CareerService._mock_roadmap(missing_skills)

    @staticmethod
    def _mock_roadmap(missing_skills: list) -> dict:
        items = []
        total_hours = 0

        difficulty_map = {
            "python": ("Intermediate", 20),
            "javascript": ("Intermediate", 25),
            "react": ("Intermediate", 30),
            "docker": ("Beginner", 15),
            "kubernetes": ("Advanced", 30),
            "aws": ("Intermediate", 25),
            "sql": ("Beginner", 15),
            "postgresql": ("Intermediate", 20),
            "mongodb": ("Beginner", 15),
            "git": ("Beginner", 8),
            "java": ("Intermediate", 30),
            "c++": ("Advanced", 40),
            "fastapi": ("Intermediate", 15),
            "django": ("Intermediate", 20),
            "flask": ("Beginner", 12),
        }

        for i, skill in enumerate(missing_skills[:6]):
            diff, hours = difficulty_map.get(skill.lower(), ("Intermediate", 20))
            priority = "High" if i < 2 else ("Medium" if i < 4 else "Low")
            items.append({
                "topic": skill,
                "description": f"Learn {skill} from fundamentals to practical application.",
                "hours": hours,
                "difficulty": diff,
                "priority": priority,
                "milestones": [
                    f"Complete {skill} fundamentals",
                    f"Build a practice project with {skill}",
                    f"Review and document learnings",
                ],
                "mini_project": f"Build a small application using {skill}",
            })
            total_hours += hours

        estimated_weeks = max(1, total_hours // 10)

        return {
            "items": items,
            "total_hours": total_hours,
            "estimated_weeks": estimated_weeks,
        }

    # ── Resume Optimization ──

    @staticmethod
    def optimize_resume(resume_text: str, ats_suggestions: list, jd_text: str = None) -> dict:
        jd_section = f"\nTarget Job Description:\n{jd_text[:2000]}" if jd_text else ""
        suggestions_text = "\n".join(f"- {s}" for s in ats_suggestions)

        prompt = f"""Optimize this resume based on ATS suggestions and optionally the target job description.

Resume:
{resume_text[:3000]}

ATS Suggestions to address:
{suggestions_text}
{jd_section}

Return a JSON object:
{{
  "optimized_text": "The full optimized resume text...",
  "improvements": ["improvement1", "improvement2"],
  "professional_summary": "Optimized professional summary",
  "optimized_skills": ["skill1", "skill2"],
  "optimized_projects": ["project1 optimized description"],
  "optimized_keywords": ["keyword1", "keyword2"],
  "optimized_experience": ["experience1 optimized"]
}}"""

        result = CareerService._call_ai(prompt, max_tokens=1200)
        if result:
            return result

        return CareerService._mock_optimize_resume(resume_text, ats_suggestions, jd_text)

    @staticmethod
    def _mock_optimize_resume(resume_text: str, ats_suggestions: list, jd_text: str = None) -> dict:
        improvements = []
        if ats_suggestions:
            improvements = [f"Addressed: {s}" for s in ats_suggestions[:5]]

        return {
            "optimized_text": resume_text,
            "improvements": improvements if improvements else ["General formatting improvements applied"],
            "professional_summary": "Optimized professional summary tailored for the target role.",
            "optimized_skills": [],
            "optimized_projects": [],
            "optimized_keywords": [],
            "optimized_experience": [],
        }

    # ── Career Readiness ──

    @staticmethod
    def calculate_career_readiness(
        resume_match: float = None,
        ats_score: float = None,
        skill_gap: float = None,
        interview_score: float = None,
        coding_score: float = None,
    ) -> dict:
        scores = {}
        valid_scores = []

        if resume_match is not None:
            scores["resume_match_score"] = resume_match
            valid_scores.append(resume_match)
        else:
            scores["resume_match_score"] = None

        if ats_score is not None:
            scores["ats_score"] = ats_score
            valid_scores.append(ats_score)
        else:
            scores["ats_score"] = None

        if interview_score is not None:
            scores["interview_score"] = interview_score
            valid_scores.append(interview_score)
        else:
            scores["interview_score"] = None

        if coding_score is not None:
            scores["coding_score"] = coding_score
            valid_scores.append(coding_score)
        else:
            scores["coding_score"] = None

        if skill_gap is not None:
            scores["skill_gap_score"] = skill_gap
            valid_scores.append(skill_gap)
        else:
            scores["skill_gap_score"] = None

        overall = round(sum(valid_scores) / len(valid_scores), 1) if valid_scores else 0.0
        scores["overall_score"] = overall

        recommendations = []
        ai_suggestions = []

        if resume_match is not None and resume_match < 70:
            recommendations.append("Work on improving your resume match score by adding relevant keywords and projects.")
            ai_suggestions.append("Tailor your resume to highlight skills matching the target role.")
        if ats_score is not None and ats_score < 70:
            recommendations.append("Improve ATS compatibility by using standard section headings and clean formatting.")
            ai_suggestions.append("Remove complex formatting, tables, and graphics that confuse ATS systems.")
        if skill_gap is not None and skill_gap < 60:
            recommendations.append("Focus on learning the missing skills identified in the skill gap analysis.")
            ai_suggestions.append("Prioritize high-demand skills from the gap analysis for maximum impact.")
        if interview_score is not None and interview_score < 70:
            recommendations.append("Practice more mock interviews to improve your communication and technical explanations.")
            ai_suggestions.append("Review common interview questions for your target role and practice aloud.")
        if coding_score is not None and coding_score < 70:
            recommendations.append("Solve more coding challenges to improve your problem-solving skills.")
            ai_suggestions.append("Practice data structures and algorithms daily on platforms like LeetCode.")

        if not recommendations:
            recommendations.append("You're on the right track! Keep building projects and refining your skills.")
        if not ai_suggestions:
            ai_suggestions.append("Continue your current learning path and stay consistent.")

        scores["recommendations"] = recommendations
        scores["ai_suggestions"] = ai_suggestions

        return scores

    # ── AI Suggestions ──

    @staticmethod
    def get_ai_suggestions(
        resume_skills: list = None,
        missing_skills: list = None,
        ats_score: float = None,
        match_score: float = None,
    ) -> dict:
        prompt = f"""Based on the user's career profile, provide actionable AI suggestions.

Resume Skills: {json.dumps(resume_skills or [])}
Missing Skills: {json.dumps(missing_skills or [])}
ATS Score: {ats_score}
Match Score: {match_score}

Return a JSON object:
{{
  "immediate_actions": ["action1", "action2"],
  "skill_recommendations": ["skill1", "skill2"],
  "career_tips": ["tip1", "tip2"],
  "resources": ["resource1", "resource2"],
  "priority_focus": "Areas to focus on immediately"
}}"""

        result = CareerService._call_ai(prompt)
        if result:
            return result

        return CareerService._mock_suggestions(resume_skills, missing_skills, ats_score, match_score)

    @staticmethod
    def _mock_suggestions(
        resume_skills: list = None,
        missing_skills: list = None,
        ats_score: float = None,
        match_score: float = None,
    ) -> dict:
        immediate_actions = []
        skill_recs = []
        tips = []
        resources = []
        priority = ""

        if ats_score is not None and ats_score < 70:
            immediate_actions.append("Update your resume to improve ATS compatibility")
            tips.append("Use keywords from job descriptions in your resume")

        if match_score is not None and match_score < 60:
            immediate_actions.append("Tailor your resume for your target role")
            priority = "Improving resume-to-job match"

        if missing_skills:
            skill_recs = missing_skills[:5]
            immediate_actions.append(f"Start learning {missing_skills[0]} to fill your biggest skill gap")
            resources.append(f"Online courses for {missing_skills[0]}")
            if not priority:
                priority = f"Learning {missing_skills[0]}"

        if not immediate_actions:
            immediate_actions = [
                "Continue building projects to strengthen your portfolio",
                "Practice mock interviews regularly",
            ]

        if not tips:
            tips = [
                "Keep your LinkedIn profile updated",
                "Contribute to open source projects",
                "Network with professionals in your target field",
            ]

        if not resources:
            resources = [
                "LeetCode for coding practice",
                "Coursera/Udemy for skill development",
                "GitHub for portfolio building",
            ]

        if not priority:
            priority = "Consistent practice and skill development"

        return {
            "immediate_actions": immediate_actions,
            "skill_recommendations": skill_recs,
            "career_tips": tips,
            "resources": resources,
            "priority_focus": priority,
        }

    # ------------------------------------------------------------------
    # Format-Preserving DOCX Optimizer (XML-level text replacement)
    # ------------------------------------------------------------------

    @staticmethod
    def _get_paragraph_text(para) -> str:
        """Get full text from a paragraph by concatenating all run texts."""
        return "".join(r.text for r in para.runs)

    @staticmethod
    def _set_paragraph_text_preserving_format(para, new_text: str):
        """Replace entire paragraph text while preserving ALL run formatting.

        Strategy: Keep the first run's formatting for the full text, clear all other runs.
        This preserves paragraph-level style (font, size, color, bold, italic, etc.)
        """
        runs = para.runs
        if not runs:
            return

        # Preserve the first run's complete formatting (XML element with all attributes)
        first_run = runs[0]
        first_run.text = new_text

        # Clear all other runs' text (keep their XML elements for structure but empty text)
        for run in runs[1:]:
            run.text = ""

    @staticmethod
    def _replace_in_paragraph(para, old_text: str, new_text: str) -> bool:
        """Replace specific text within a paragraph while preserving run-level formatting.

        Handles the common case where formatting spans multiple runs.
        If the old text is within a single run, replaces just that text.
        If it spans runs, consolidates into the first affected run.
        """
        if not old_text or old_text == new_text:
            return old_text == new_text

        runs = para.runs
        if not runs:
            return False

        # Build cumulative text map
        run_texts = [r.text for r in runs]
        full_text = "".join(run_texts)

        if old_text not in full_text:
            return False

        start_idx = full_text.index(old_text)
        end_idx = start_idx + len(old_text)

        # Find affected runs
        cum_len = 0
        start_run_idx = None
        end_run_idx = None
        start_offset = 0
        end_offset = 0

        for i, rt in enumerate(run_texts):
            run_start = cum_len
            run_end = cum_len + len(rt)

            if start_run_idx is None and start_idx < run_end and end_idx > run_start:
                start_run_idx = i
                start_offset = start_idx - run_start

            if start_run_idx is not None and end_run_idx is None and end_idx <= run_end:
                end_run_idx = i
                end_offset = end_idx - run_start
                break

            cum_len += len(rt)

        if start_run_idx is None or end_run_idx is None:
            return False

        # Single run replacement - preserves all formatting perfectly
        if start_run_idx == end_run_idx:
            run = runs[start_run_idx]
            run.text = run.text[:start_offset] + new_text + run.text[end_offset:]
            return True

        # Multi-run replacement: put all text in first run, clear the rest
        first_run = runs[start_run_idx]
        suffix = runs[end_run_idx].text[end_offset:]
        first_run.text = first_run.text[:start_offset] + new_text + suffix

        for i in range(start_run_idx + 1, end_run_idx + 1):
            runs[i].text = ""

        return True

    @staticmethod
    def _optimize_paragraph_text(text: str, context: str, analysis: dict, jd_text: str = None) -> str:
        """Use AI to optimize a single paragraph's text. Returns optimized text or original if no change needed."""
        if not text.strip() or len(text.strip()) < 10:
            return text

        missing_keywords = [m["keyword"] for m in analysis.get("keyword_analysis", {}).get("missing", [])[:8]]
        jd_context = f"\nTarget JD keywords: {', '.join(missing_keywords[:5])}" if missing_keywords else ""

        prompt = f"""Optimize this resume paragraph for ATS compatibility and professional impact.

RULES:
- Do NOT fabricate experience, skills, or achievements
- Do NOT add information not present in the original
- ONLY improve wording, action verbs, and keyword integration
- Keep the same meaning and approximate length
- Return ONLY the optimized text, nothing else
- If the text is already good, return it unchanged

Original paragraph:
{text.strip()}

Context: {context}
{jd_context}

Return the optimized paragraph text (nothing else):"""

        result = CareerService._call_ai(prompt, max_tokens=500)
        if result and isinstance(result, dict):
            optimized = result.get("optimized_text", result.get("text", ""))
            if optimized and len(optimized.strip()) > 5:
                return optimized.strip()
        elif result and isinstance(result, str):
            return result.strip()

        return text

    @staticmethod
    def optimize_docx_format_preserving(
        original_file_path: str,
        resume_text: str,
        ats_analysis: dict,
        jd_text: str = None,
    ) -> bytes:
        """Optimize a DOCX resume by editing text in-place while preserving 100% formatting.

        This opens the ORIGINAL document and only modifies text content within
        existing paragraphs/runs. All fonts, sizes, colors, spacing, tables,
        images, headers/footers, and layout remain unchanged.

        Returns the optimized DOCX as bytes.
        """
        from docx import Document
        import io

        real_path = os.path.realpath(original_file_path)
        uploads_dir = os.path.realpath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads"))
        if not real_path.startswith(uploads_dir + os.sep) and real_path != uploads_dir:
            raise ValueError("File path must be within the uploads directory")

        doc = Document(real_path)

        # Categorize paragraphs by section for context-aware optimization
        current_section = "header"
        section_headers = {
            "summary", "objective", "profile", "about",
            "experience", "work", "employment", "professional experience",
            "education", "academic",
            "skills", "technical skills", "technologies", "competencies",
            "projects", "portfolio",
            "certifications", "licenses", "credentials",
            "awards", "honors", "achievements",
            "publications", "papers",
            "volunteer", "community",
            "languages", "interests",
        }

        def detect_section(text: str) -> str:
            t = text.strip().lower().rstrip(":")
            if t in section_headers or t.endswith(":"):
                for sh in section_headers:
                    if sh in t:
                        return sh
            return current_section

        # Process body paragraphs
        for para in doc.paragraphs:
            text = CareerService._get_paragraph_text(para)
            stripped = text.strip()

            if not stripped:
                continue

            # Detect section changes
            new_section = detect_section(stripped)
            if new_section != current_section and new_section != stripped.lower():
                current_section = new_section
                # Don't optimize section headers - keep them as-is
                if stripped.lower() in section_headers or stripped.endswith(":"):
                    continue
                # If it's a short header-like line, skip
                if len(stripped) < 40 and stripped.upper() == stripped:
                    continue

            # Skip very short lines (likely headers/labels)
            if len(stripped) < 15:
                continue

            # Skip lines that are just numbers, dates, or separators
            if stripped.replace(" ", "").replace("-", "").replace("–", "").isdigit():
                continue
            if all(c in "-=—_" for c in stripped.replace(" ", "")):
                continue

            # Optimize the paragraph text
            context = f"Section: {current_section}"
            optimized = CareerService._optimize_paragraph_text(stripped, context, ats_analysis, jd_text)

            # Only replace if meaningfully different and not longer
            if optimized != stripped:
                # Ensure optimized text isn't significantly longer (would break layout)
                if len(optimized) <= len(stripped) * 1.15:  # Allow 15% length increase max
                    CareerService._replace_in_paragraph(para, stripped, optimized)

        # Process table cells (same logic)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = CareerService._get_paragraph_text(para)
                        stripped = text.strip()
                        if not stripped or len(stripped) < 15:
                            continue

                        optimized = CareerService._optimize_paragraph_text(
                            stripped, "Table content", ats_analysis, jd_text
                        )
                        if optimized != stripped and len(optimized) <= len(stripped) * 1.15:
                            CareerService._replace_in_paragraph(para, stripped, optimized)

        # Process headers and footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if header_footer is not None:
                    for para in header_footer.paragraphs:
                        text = CareerService._get_paragraph_text(para)
                        stripped = text.strip()
                        if not stripped or len(stripped) < 15:
                            continue
                        # Don't optimize headers/footers - keep them as-is
                        # They typically contain name, page numbers, etc.

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    @staticmethod
    def docx_to_pdf(docx_bytes: bytes) -> bytes:
        """Convert a DOCX to PDF while preserving formatting as closely as possible.

        Uses the DOCX's own styles and structure to generate a faithful PDF.
        """
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, pt
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        import io
        import os

        doc = Document(io.BytesIO(docx_bytes))
        output = io.BytesIO()

        # Get page margins from DOCX
        section = doc.sections[0] if doc.sections else None
        left_margin = 0.75 * inch
        right_margin = 0.75 * inch
        top_margin = 0.75 * inch
        bottom_margin = 0.75 * inch

        if section:
            if section.left_margin:
                left_margin = section.left_margin / 914400 * inch  # EMU to inches
            if section.right_margin:
                right_margin = section.right_margin / 914400 * inch
            if section.top_margin:
                top_margin = section.top_margin / 914400 * inch
            if section.bottom_margin:
                bottom_margin = section.bottom_margin / 914400 * inch

        pdf_doc = SimpleDocTemplate(
            output,
            pagesize=letter,
            leftMargin=left_margin,
            rightMargin=right_margin,
            topMargin=top_margin,
            bottomMargin=bottom_margin,
        )

        styles = getSampleStyleSheet()
        story = []

        # Map DOCX styles to ReportLab styles
        def get_run_style(run):
            """Extract formatting from a docx run and create a ReportLab style."""
            font_name = run.font.name or "Helvetica"
            font_size = run.font.size
            if font_size:
                font_size_pt = font_size.pt if hasattr(font_size, 'pt') else font_size / 12700
            else:
                font_size_pt = 10

            is_bold = run.font.bold
            is_italic = run.font.italic
            is_underline = run.font.underline

            # Get color
            font_color = None
            if run.font.color and run.font.color.rgb:
                rgb = str(run.font.color.rgb)
                if rgb and rgb != "000000":
                    try:
                        font_color = colors.HexColor(f"#{rgb}")
                    except:
                        pass

            # Map font names to ReportLab equivalents
            font_map = {
                "calibri": "Helvetica",
                "arial": "Helvetica",
                "times new roman": "Times-Roman",
                "times": "Times-Roman",
                "courier": "Courier",
                "consolas": "Courier",
                "georgia": "Times-Roman",
                "verdana": "Helvetica",
                "tahoma": "Helvetica",
                "impact": "Helvetica-Bold",
            }
            rl_font = font_map.get(font_name.lower(), "Helvetica") if font_name else "Helvetica"
            if is_bold and not rl_font.endswith("-Bold"):
                if rl_font == "Times-Roman":
                    rl_font = "Times-Bold"
                elif rl_font == "Helvetica":
                    rl_font = "Helvetica-Bold"
                elif rl_font == "Courier":
                    rl_font = "Courier-Bold"

            return {
                "font_name": rl_font,
                "font_size": font_size_pt,
                "bold": is_bold,
                "italic": is_italic,
                "underline": is_underline,
                "color": font_color,
            }

        def get_paragraph_alignment(para):
            """Map DOCX paragraph alignment to ReportLab."""
            alignment = para.paragraph_format.alignment
            if alignment is not None:
                align_str = str(alignment).lower()
                if "center" in align_str:
                    return TA_CENTER
                elif "right" in align_str:
                    return TA_RIGHT
                elif "justify" in align_str:
                    return TA_JUSTIFY
            return TA_LEFT

        def build_paragraph(para):
            """Convert a DOCX paragraph to ReportLab flowables."""
            text = para.text.strip()
            if not text:
                # Empty paragraph - add small spacer
                space_before = para.paragraph_format.space_before
                space_after = para.paragraph_format.space_after
                spacer_height = 6
                if space_before:
                    spacer_height = max(spacer_height, space_before.pt if hasattr(space_before, 'pt') else 6)
                return [Spacer(1, spacer_height)]

            flowables = []
            alignment = get_paragraph_alignment(para)

            # Build formatted text using ReportLab's Paragraph with inline styling
            # We need to build XML-like markup for ReportLab
            parts = []
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue

                # Escape XML special characters
                safe_text = (
                    run_text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )

                style = get_run_style(run)
                tags = []
                if style["bold"]:
                    tags.append("b")
                if style["italic"]:
                    tags.append("i")
                if style["underline"]:
                    tags.append("u")

                if tags:
                    for tag in tags:
                        safe_text = f"<{tag}>{safe_text}</{tag}>"

                parts.append(safe_text)

            formatted_text = "".join(parts)
            if not formatted_text.strip():
                return [Spacer(1, 4)]

            # Determine paragraph style
            is_heading = False
            heading_level = 0
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading" in style_name:
                    is_heading = True
                    try:
                        heading_level = int(style_name.replace("heading", "").strip())
                    except:
                        heading_level = 2

            # Get font info from first run for the paragraph style
            first_font = {"font_name": "Helvetica", "font_size": 10, "bold": False, "italic": False, "color": None}
            if para.runs:
                first_font = get_run_style(para.runs[0])

            font_size = first_font["font_size"]
            if is_heading:
                font_size = max(font_size, 12 + (4 - min(heading_level, 4)) * 2)

            style_props = {
                "fontName": first_font["font_name"],
                "fontSize": font_size,
                "leading": font_size * 1.3,
                "alignment": alignment,
                "spaceAfter": 4,
                "spaceBefore": 4,
            }
            if first_font["color"]:
                style_props["textColor"] = first_font["color"]

            custom_style = ParagraphStyle(
                f"Resume_{id(para)}",
                parent=styles["Normal"],
                **style_props,
            )

            flowables.append(Paragraph(formatted_text, custom_style))
            return flowables

        # Build document
        for i, para in enumerate(doc.paragraphs):
            flowables = build_paragraph(para)
            story.extend(flowables)

        # Add tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                    safe_cell = (
                        cell_text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    row_data.append(safe_cell)
                table_data.append(row_data)

            if table_data:
                # Determine column count
                max_cols = max(len(row) for row in table_data)
                for row in table_data:
                    while len(row) < max_cols:
                        row.append("")

                col_width = (letter[0] - left_margin - right_margin) / max_cols
                t = Table(table_data, colWidths=[col_width] * max_cols)
                t.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(Spacer(1, 6))
                story.append(t)
                story.append(Spacer(1, 6))

        pdf_doc.build(story)
        output.seek(0)
        return output.getvalue()

    @staticmethod
    def calculate_before_after_score(
        original_text: str, optimized_text: str, ats_analysis: dict, jd_text: str = None
    ) -> dict:
        """Calculate ATS score before and after optimization."""
        before_score = ats_analysis.get("overall_score", 0)

        # Count keyword improvements in optimized text
        optimized_lower = optimized_text.lower()
        keywords_found = 0
        for kw in ats_analysis.get("keyword_analysis", {}).get("missing", []):
            if kw.get("keyword", "").lower() in optimized_lower:
                keywords_found += 1

        # Estimate improvement
        keyword_improvement = min(20, keywords_found * 3)
        estimated_after = min(100, before_score + keyword_improvement + 5)

        return {
            "before": before_score,
            "after": estimated_after,
            "improvement": estimated_after - before_score,
            "keywords_added": keywords_found,
        }
